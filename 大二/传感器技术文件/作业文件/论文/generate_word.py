#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_word.py — 自动驾驶汽车先进传感器技术综述 Word 文档生成器
遵循广东工业大学(GDUT)毕业论文格式规范

运行: D:/program/anaconda3/python.exe generate_word.py
输出: D:/sensorhomework/论文/自动驾驶汽车先进传感器技术综述.docx

Formatting (GDUT Standard):
  - A4, margins top/bottom=2.54cm, left/right=3.17cm
  - Title: SimHei 22pt bold centered
  - Abstract title: SimHei 16pt bold
  - Abstract body: KaiTi 10.5pt
  - Keywords: KaiTi 10.5pt
  - Chapter (一级): SimHei 16pt bold
  - Section (二级): SimHei 12pt bold (小四号)
  - Subsection (三级): SimHei 12pt bold (小四号)
  - Body text: SimSun 12pt (小四号), 1.5× line spacing
  - References: SimSun 12pt (小四号)
  - Page numbers centered at bottom
"""

import os
import sys

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Constants ──────────────────────────────────────────────────────────
FIGURE_DIR = r"D:\sensorhomework\论文\figure"

COVER_TITLE_TEXT = "自动驾驶汽车先进传感器技术综述：\n原理、测量电路与多模态融合应用"

# ── Helper Functions ────────────────────────────────────────────────────

def set_cell_font(cell, font_name, font_name_east, size_pt, bold=False, color=None):
    """Set font for all runs in a table cell."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_east)
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color


def add_run(paragraph, text, font_name, font_name_east, size_pt, bold=False,
            italic=False, color=None, underline=False):
    """Add a formatted run to a paragraph and return the run."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_east)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    if underline:
        run.font.underline = True
    return run


def new_paragraph(doc, text="", style_name='Normal', alignment=None,
                  space_before=0, space_after=0, first_line_indent=None,
                  line_spacing=1.5):
    """Create a new paragraph with standard formatting."""
    p = doc.add_paragraph(style=style_name)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    return p


def _set_outline_level(paragraph, level):
    """Set the outline level on a paragraph so TOC field can find it."""
    pPr = paragraph._p.get_or_add_pPr()
    # Remove existing outlineLvl if any
    existing = pPr.findall(qn('w:outlineLvl'))
    for e in existing:
        pPr.remove(e)
    outline_lvl = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{level}"/>')
    pPr.append(outline_lvl)


def heading_chapter(doc, text):
    """Add a chapter heading (一级标题): SimHei 16pt bold, centered. 三号黑体加粗"""
    p = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=18, space_after=12, line_spacing=1.5)
    add_run(p, text, 'SimHei', 'SimHei', 16, bold=True)
    _set_outline_level(p, 0)
    return p


def heading_section(doc, text):
    """Add a section heading (二级标题): SimHei 14pt bold, left-aligned. 四号黑体加粗"""
    p = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      space_before=12, space_after=6, line_spacing=1.5)
    add_run(p, text, 'SimHei', 'SimHei', 14, bold=True)
    _set_outline_level(p, 1)
    return p


def heading_subsection(doc, text):
    """Add a subsection heading (三级标题): SimHei 12pt bold, left-aligned. 小四号黑体加粗"""
    p = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      space_before=8, space_after=4, line_spacing=1.5)
    add_run(p, text, 'SimHei', 'SimHei', 12, bold=True)
    _set_outline_level(p, 2)
    return p


def body_paragraph(doc, text, first_line_indent=Cm(0.74)):
    """Add body text: SimSun 12pt, 1.5× line spacing, with first-line indent."""
    p = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      space_before=0, space_after=0,
                      first_line_indent=first_line_indent,
                      line_spacing=1.5)
    add_run(p, text, 'SimSun', 'SimSun', 12, bold=False)
    return p


def body_paragraph_no_indent(doc, text):
    """Body text without first-line indent."""
    return body_paragraph(doc, text, first_line_indent=None)


def insert_image(doc, image_name, width_inches=5.5, caption="", alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Insert an image from the figure directory with optional caption."""
    img_path = os.path.join(FIGURE_DIR, image_name)
    if os.path.exists(img_path):
        p_img = new_paragraph(doc, alignment=alignment,
                              space_before=6, space_after=3, line_spacing=1.0)
        run = p_img.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
        if caption:
            p_cap = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  space_before=3, space_after=6, line_spacing=1.2)
            add_run(p_cap, caption, 'SimSun', 'SimSun', 9, bold=False)
    else:
        p_note = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               space_before=6, space_after=6, line_spacing=1.0)
        add_run(p_note, f"[图 {image_name} 未找到，请放入 figure/ 目录]",
                'SimSun', 'SimSun', 9, bold=False, color=RGBColor(128, 128, 128))
    return


FORMULA_DIR = r"D:\sensorhomework\论文\formulas\png"
CIRCUIT_DIR = r"D:\sensorhomework\论文\circuits"

def add_circuit_img(doc, circuit_name, width_inches=5.5, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Insert a circuit diagram PNG with caption"""
    img_path = os.path.join(CIRCUIT_DIR, f"{circuit_name}.png")
    if os.path.exists(img_path):
        new_paragraph(doc, "", space_before=6, space_after=0, line_spacing=1.0)
        p = new_paragraph(doc, alignment=alignment, space_before=4, space_after=4)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))

def add_formula_img(doc, formula_name, width_inches=4.5, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Insert a rendered LaTeX formula PNG as an inline image."""
    img_path = os.path.join(FORMULA_DIR, f"{formula_name}.png")
    if os.path.exists(img_path):
        p = new_paragraph(doc, alignment=alignment, space_before=4, space_after=4)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
    else:
        p = new_paragraph(doc, alignment=alignment)
        add_run(p, f"[公式 {formula_name} 未渲染，请运行公式生成脚本]",
                'SimSun', 'SimSun', 9, color=RGBColor(192, 0, 0))

def key_insight_box(doc, text):
    """Add a highlighted key insight box (用边框段落模拟)."""
    p = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      space_before=8, space_after=8, line_spacing=1.5)
    # Add shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF3E0" w:val="clear"/>')
    pPr.append(shd)
    # Add border
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="4" w:color="E86A17"/>'
        f'<w:left w:val="single" w:sz="4" w:space="4" w:color="E86A17"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="E86A17"/>'
        f'<w:right w:val="single" w:sz="4" w:space="4" w:color="E86A17"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    add_run(p, "[核心洞察] ", 'SimHei', 'SimHei', 10.5, bold=True, color=RGBColor(0xE8, 0x6A, 0x17))
    add_run(p, text, 'SimSun', 'SimSun', 10.5, bold=False)
    return p


def add_table(doc, headers, rows, col_widths=None, caption="", font_size=9):
    """Add a formatted table with header and data rows."""
    if caption:
        p_cap = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              space_before=6, space_after=3, line_spacing=1.2)
        add_run(p_cap, caption, 'SimHei', 'SimHei', 9, bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, 'SimHei', 'SimHei', font_size, bold=True)
        # Shade header
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0F2448" w:val="clear"/>'
        )
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        # Header text white
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r, row_data in enumerate(rows):
        row_cells = table.rows[r + 1].cells
        for c, val in enumerate(row_data):
            row_cells[c].text = ""
            p = row_cells[c].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, str(val), 'SimSun', 'SimSun', font_size, bold=False)
            # Alternate row shading
            if r % 2 == 1:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="F2F7FB" w:val="clear"/>'
                )
                row_cells[c]._tc.get_or_add_tcPr().append(shading)

    # Set column widths if specified
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    # Add spacing after table
    new_paragraph(doc, "", space_before=3, space_after=3, line_spacing=1.0)
    return table


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run._element.append(parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>'))
    return p


def add_section_break(doc):
    """Add a section break (new page)."""
    new_section = doc.add_section()
    return new_section


# ── Main Document Generation ────────────────────────────────────────────

def create_document():
    """Create the complete Word document."""

    doc = Document()

    # ── Page Setup (A4) ─────────────────────────────────────────────
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # Ensure subsequent sections keep the same margins
    # We'll configure default section properties

    # ── Style configuration ─────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # ── Page Numbers (footer, centered) ─────────────────────────────
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add PAGE field
        run = fp.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        run2 = fp.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)
        run3 = fp.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)
        for r in [run, run2, run3]:
            r.font.name = 'SimSun'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            r.font.size = Pt(9)

    # =====================================================================
    # 1. COVER PAGE
    # =====================================================================
    # Add empty paragraphs for vertical spacing
    for _ in range(6):
        new_paragraph(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=0, space_after=0, line_spacing=1.0)

    p = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=0, space_after=18, line_spacing=1.0)
    add_run(p, "广东工业大学", 'SimHei', 'SimHei', 22, bold=True)

    new_paragraph(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=0, space_after=0, line_spacing=1.0)

    p = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=24, space_after=24, line_spacing=1.2)
    add_run(p, "传感器技术与应用  调研报告", 'SimHei', 'SimHei', 16, bold=True)

    # Title
    p = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=24, space_after=6, line_spacing=1.5)
    add_run(p, "论文题目：自动驾驶汽车先进传感器技术综述：", 'SimHei', 'SimHei', 15, bold=True)

    p = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=0, space_after=24, line_spacing=1.5)
    add_run(p, "原理、测量电路与多模态融合应用", 'SimHei', 'SimHei', 15, bold=True)

    # Info fields
    info_items = ["学院", "专业", "学号", "姓名", "指导教师"]
    for item in info_items:
        p = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_before=6, space_after=6, line_spacing=2.0)
        if item == "指导教师":
            add_run(p, f"{item}：翟老师（zhaiyh@gdut.edu.cn）", 'SimSun', 'SimSun', 14, bold=False)
        else:
            add_run(p, f"{item}：_______________", 'SimSun', 'SimSun', 14, bold=False)

    # Date
    new_paragraph(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=36, space_after=0, line_spacing=1.0)
    p = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=6, space_after=0, line_spacing=1.0)
    add_run(p, "2026 年 5 月", 'SimSun', 'SimSun', 14, bold=False)

    add_page_break(doc)

    # =====================================================================
    # 2. CHINESE ABSTRACT (样张2: 摘要三号黑体加粗, 正文小四宋体, 关键词四号黑体+五号宋体)
    # =====================================================================
    heading_chapter(doc, "摘    要")

    abstract_body = (
        "自动驾驶技术正在深刻变革全球汽车产业格局。作为智能网联汽车的「感知器官」，先进传感器系统是实现"
        "L3及以上级别自动驾驶的核心基础。本文面向自动驾驶汽车环境感知需求，对五大类车载传感器——视觉摄像机、"
        "毫米波雷达、超声波雷达、激光雷达（LiDAR）及惯性测量单元（IMU）——进行了系统性综述。"
    )
    body_paragraph(doc, abstract_body)

    abstract_body2 = (
        "在物理原理层面，本文从光电效应、电磁波传播、压电效应及牛顿惯性定律出发，分别推导了CMOS图像传感器的"
        "光电转换模型、FMCW毫米波雷达的测距测速公式、超声波雷达的渡越时间方程、激光雷达的ToF与FMCW数学模型，"
        "以及MEMS-IMU的加速度-角速度解算原理。在核心元器件层面，详细剖析了CMOS有源像素（4T-APS）、77GHz MMIC"
        "收发芯片、VCSEL/EEL发射阵列、APD/SPAD单光子探测器以及MEMS陀螺仪-加速度计等关键器件的微观工作机制。"
    )
    body_paragraph(doc, abstract_body2)

    abstract_body3 = (
        "在测量电路层面，深入分析了相关双采样（CDS）降噪电路、毫米波雷达MMIC混频-中频链路、跨阻放大器（TIA）"
        "微弱信号调理电路、时间数字转换器（TDC）皮秒级计时电路等核心信号链的设计逻辑。在系统集成层面，阐述了"
        "固态激光雷达SIP系统级封装对信号完整性的提升作用，以及SPAD淬灭-复位电路对光子计数性能的影响。"
    )
    body_paragraph(doc, abstract_body3)

    abstract_body4 = (
        "在应用与展望层面，本文围绕多传感器融合这一核心范式，系统论述了数据级、特征级与决策级融合的数学框架，"
        "并深入探讨了基于前融合的BEV（鸟瞰图）感知架构与Transformer多模态对齐算法。结合全球及中国自动驾驶传感器"
        "市场最新数据（2023-2030年），分析了激光雷达、毫米波雷达和车载摄像头三大细分市场的规模与发展趋势，"
        "并基于产业链核心企业的招聘信息进行了人才需求解读与职业规划建议，为电子信息类专业学生的学术研究"
        "与职业发展提供全面参考。"
    )
    body_paragraph(doc, abstract_body4)

    # Keywords (GDUT: "关键词"四号黑体加粗, 关键词内容五号宋体)
    kw_para = new_paragraph(doc, space_before=14, space_after=6, line_spacing=1.5)
    add_run(kw_para, "关键词：", 'SimHei', 'SimHei', 14, bold=True)
    add_run(kw_para, "自动驾驶；CMOS图像传感器；FMCW毫米波雷达；固态激光雷达；MEMS惯性传感器；"
            "多传感器融合；跨阻放大器(TIA)；深度学习",
            'SimSun', 'SimSun', 10.5, bold=False)

    add_page_break(doc)

    # =====================================================================
    # 3. ENGLISH ABSTRACT (样张3: Abstract三号TNR加粗, 正文小四TNR, Key words四号TNR加粗)
    # =====================================================================
    heading_chapter(doc, "ABSTRACT")

    en_abstract1 = (
        "Autonomous driving technology is profoundly transforming the global automotive industry. "
        "As the \"perceptual organs\" of intelligent connected vehicles, advanced sensor systems constitute "
        "the core foundation for achieving L3 and above autonomous driving. This report presents a systematic "
        "survey of five major categories of automotive sensors — vision cameras, millimeter-wave radars, "
        "ultrasonic radars, LiDAR, and inertial measurement units (IMU) — oriented toward environmental "
        "perception requirements for autonomous vehicles."
    )
    # English body: TNR 12pt (小四)
    p_en1 = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=0, space_after=0, first_line_indent=Cm(0.74), line_spacing=1.5)
    add_run(p_en1, en_abstract1, 'Times New Roman', 'Times New Roman', 12, bold=False)

    en_abstract2 = (
        "At the physical principle level, starting from the photoelectric effect, electromagnetic wave "
        "propagation, piezoelectric effect, and Newton's laws of inertia, this report derives the photoelectric "
        "conversion model of CMOS image sensors, the ranging and velocity measurement formulas of FMCW "
        "millimeter-wave radars, the time-of-flight equation of ultrasonic radars, the ToF and FMCW mathematical "
        "models of LiDAR, and the acceleration/angular velocity resolution principles of MEMS-IMU. At the core "
        "component level, the micro-mechanisms of CMOS active pixels (4T-APS), 77GHz MMIC transceiver chips, "
        "VCSEL/EEL emitter arrays, APD/SPAD single-photon detectors, and MEMS gyroscope-accelerometer assemblies "
        "are analyzed in detail."
    )
    p_en2 = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=0, space_after=0, first_line_indent=Cm(0.74), line_spacing=1.5)
    add_run(p_en2, en_abstract2, 'Times New Roman', 'Times New Roman', 12, bold=False)

    en_abstract3 = (
        "At the measurement circuit level, the design logic of correlated double sampling (CDS) noise reduction "
        "circuits, millimeter-wave radar MMIC mixer-IF chains, transimpedance amplifier (TIA) weak signal "
        "conditioning circuits, and time-to-digital converter (TDC) picosecond-level timing circuits are "
        "thoroughly examined. At the application and outlook level, centered on the core paradigm of multi-sensor "
        "fusion, this report systematically discusses the mathematical frameworks of data-level, feature-level, "
        "and decision-level fusion, and deeply explores the BEV (Bird's Eye View) perception architecture based "
        "on early fusion and Transformer-based multimodal alignment algorithms. Furthermore, combined with the "
        "latest global and Chinese autonomous driving sensor market data (2023-2030), market prospect analysis "
        "and talent demand interpretation for typical positions are provided, offering references for career "
        "planning for students majoring in electronic information."
    )
    p_en3 = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=0, space_after=0, first_line_indent=Cm(0.74), line_spacing=1.5)
    add_run(p_en3, en_abstract3, 'Times New Roman', 'Times New Roman', 12, bold=False)

    # Key words (GDUT: "Key words"四号TNR加粗)
    en_kw_para = new_paragraph(doc, space_before=14, space_after=6, line_spacing=1.5)
    add_run(en_kw_para, "Key words: ", 'Times New Roman', 'Times New Roman', 14, bold=True)
    add_run(en_kw_para,
            "Autonomous Driving; CMOS Image Sensor; FMCW Millimeter-Wave Radar; Solid-State LiDAR; "
            "MEMS Inertial Sensor; Multi-Sensor Fusion; Transimpedance Amplifier (TIA); Deep Learning",
            'Times New Roman', 'Times New Roman', 12, bold=False)

    add_page_break(doc)

    # =====================================================================
    # 4. TABLE OF CONTENTS (样张4: 目录三号黑体加粗)
    # =====================================================================
    heading_chapter(doc, "目　录")

    # Build TOC programmatically - all headings known from script structure
    toc_entries = [
        (0, "摘要"),
        (0, "ABSTRACT"),
        (0, "第一章  绪论与市场前景分析"),
        (1, "1.1  研究背景：自动驾驶的感知需求"),
        (1, "1.2  全球与中国自动驾驶传感器市场规模"),
        (1, "1.3  核心传感器细分市场分析"),
        (2, "1.3.1  激光雷达市场"),
        (2, "1.3.2  毫米波雷达市场"),
        (2, "1.3.3  车载摄像头市场"),
        (1, "1.4  典型企业及岗位分析"),
        (2, "1.4.1  其域创新（XGRIDS）"),
        (2, "1.4.2  禾赛科技与速腾聚创"),
        (0, "第二章  视觉摄像机传感器"),
        (1, "2.1  工作原理"),
        (2, "2.1.1  光电转换的物理基础"),
        (2, "2.1.2  CMOS有源像素结构（4T-APS）"),
        (1, "2.2  核心元器件"),
        (2, "2.2.1  CMOS图像传感器芯片架构"),
        (2, "2.2.2  Pinned光电二极管与微透镜阵列"),
        (1, "2.3  测量电路"),
        (2, "2.3.1  相关双采样（CDS）降噪电路"),
        (2, "2.3.2  列级ADC架构"),
        (1, "2.4  应用场景"),
        (0, "第三章  毫米波雷达传感器"),
        (1, "3.1  工作原理"),
        (2, "3.1.1  FMCW信号模型"),
        (2, "3.1.2  混频与中频信号处理"),
        (2, "3.1.3  角度估计与4D成像雷达"),
        (1, "3.2  核心元器件"),
        (2, "3.2.1  77GHz MMIC收发芯片"),
        (2, "3.2.2  MIMO天线阵列"),
        (1, "3.3  测量电路"),
        (2, "3.3.1  PLL频率合成器与Chirp生成"),
        (2, "3.3.2  混频器与中频链路"),
        (1, "3.4  应用场景"),
        (0, "第四章  超声波雷达传感器"),
        (1, "4.1  工作原理"),
        (2, "4.1.1  压电效应与换能机制"),
        (2, "4.1.2  渡越时间测距模型"),
        (2, "4.1.3  温度补偿"),
        (1, "4.2  核心元器件"),
        (2, "4.2.1  压电陶瓷换能器"),
        (2, "4.2.2  ASIC集成驱动芯片"),
        (1, "4.3  测量电路"),
        (2, "4.3.1  发射驱动电路"),
        (2, "4.3.2  回波接收与阈值检测链路"),
        (1, "4.4  应用场景"),
        (0, "第五章  激光雷达（LiDAR）传感器"),
        (1, "5.1  工作原理"),
        (2, "5.1.1  内光电效应与单光子探测"),
        (2, "5.1.2  飞行时间法（ToF）测距模型"),
        (2, "5.1.3  调频连续波（FMCW）激光雷达"),
        (1, "5.2  核心元器件"),
        (2, "5.2.1  发射端：EEL、VCSEL与多结VCSEL"),
        (2, "5.2.2  接收端：PIN、APD、SPAD与SiPM"),
        (2, "5.2.3  固态扫描器件：MEMS微振镜与OPA"),
        (1, "5.3  测量电路"),
        (2, "5.3.1  跨阻放大器（TIA）"),
        (2, "5.3.2  时间数字转换器（TDC）"),
        (2, "5.3.3  淬灭与复位电路"),
        (2, "5.3.4  暗电流与SIP系统级封装"),
        (1, "5.4  应用场景"),
        (0, "第六章  惯性测量单元（IMU）传感器"),
        (1, "6.1  工作原理"),
        (2, "6.1.1  MEMS加速度计"),
        (2, "6.1.2  MEMS陀螺仪与科里奥利力效应"),
        (1, "6.2  核心元器件"),
        (2, "6.2.1  梳齿电容式MEMS加速度计"),
        (2, "6.2.2  MEMS振动陀螺仪"),
        (1, "6.3  测量电路"),
        (2, "6.3.1  开关电容读出电路"),
        (2, "6.3.2  Allan方差与噪声辨识"),
        (1, "6.4  应用场景"),
        (0, "第七章  多传感器融合与深度学习应用"),
        (1, "7.1  多传感器融合的必要性"),
        (1, "7.2  传感器融合的三个层级"),
        (2, "7.2.1  数据级融合（Early Fusion）"),
        (2, "7.2.2  特征级融合（Feature Fusion）"),
        (2, "7.2.3  决策级融合（Late Fusion）"),
        (1, "7.3  BEV感知架构与Transformer融合"),
        (1, "7.4  点云目标检测：PointPillars"),
        (1, "7.5  车载边缘计算部署"),
        (0, "第八章  市场前景与职业规划"),
        (1, "8.1  传感器产业的关键技术趋势"),
        (1, "8.2  传感器成本下降曲线"),
        (1, "8.3  典型岗位与技能需求"),
        (1, "8.4  薪资水平与发展前景"),
        (0, "第九章  结论与展望"),
        (0, "参考文献"),
        (0, "致谢"),
    ]

    for level, title in toc_entries:
        p = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          space_before=2, space_after=2, line_spacing=1.5,
                          first_line_indent=None)
        indent = level * Cm(1.5)
        p.paragraph_format.left_indent = indent

        if level == 0:
            add_run(p, title, 'SimHei', 'SimHei', 12, bold=True)
        elif level == 1:
            add_run(p, title, 'SimSun', 'SimSun', 12, bold=False)
        else:
            add_run(p, title, 'SimSun', 'SimSun', 12, bold=False)

    add_page_break(doc)
    print("  [TOC] Programmatic table of contents generated (%d entries)" % len(toc_entries))

    # =====================================================================
    # 4. CHAPTER 1: 绪论与市场前景分析
    # =====================================================================
    heading_chapter(doc, "第一章  绪论与市场前景分析")

    heading_section(doc, "1.1  研究背景：自动驾驶的感知需求")
    body_paragraph(doc,
        "自动驾驶技术是实现智能交通系统、降低交通事故率、提升出行效率的关键路径。根据国际自动机工程师学会"
        "（SAE）的定义，自动驾驶分为L0至L5六个等级。L2级（部分自动驾驶）仅需摄像头与毫米波雷达组合即可实现"
        "基本的自适应巡航与车道保持功能；然而，当系统向L3级（有条件自动驾驶）及以上演进时，车辆必须能够在"
        "复杂交通场景中独立完成环境感知、决策规划与控制执行，这对感知系统的精度、鲁棒性和冗余度提出了质的"
        "飞跃要求。"
    )
    body_paragraph(doc,
        "单一类型的传感器在功能上存在固有的物理局限性。视觉摄像机（CCD/CMOS）纹理与颜色信息丰富，是交通"
        "标志识别、车道线检测的基础，但在强光直射、逆光、夜间暗光及雨雾天气下性能急剧退化，且缺乏物理深度"
        "信息，单目测距误差随距离呈二次方增长。毫米波雷达对运动目标的距离与速度测量精度高，全天候工作能力"
        "强，但对静止物体的检测能力弱，角度分辨率有限，易产生金属隧道等环境下的多径反射假阳性。超声波雷达"
        "近距离探测成本极低，是泊车辅助的标准方案，但探测距离短、方向性差，无法满足高速行驶的感知需求。"
    )
    body_paragraph(doc,
        "激光雷达（LiDAR）能够直接获取高精度三维点云几何信息，抗环境光干扰能力强，但成本较高，且在雨雪雾"
        "天气中存在一定的衰减。惯性测量单元（IMU）提供载体自身运动状态的高频估计，但存在漂移累积误差，无法"
        "独立提供绝对位置信息。因此，现代L3+自动驾驶系统普遍采用多传感器融合（Multi-Sensor Fusion, MSF）"
        "架构，通过组合异质传感器以形成物理原理上的感知冗余与信息互补。这一范式是本文展开各传感器技术调研"
        "的根本出发点。"
    )

    heading_section(doc, "1.2  全球与中国自动驾驶传感器市场规模")
    body_paragraph(doc,
        "根据Gartner、Yole Intelligence、高工智能汽车研究院等多家权威市场研究机构的数据综合，全球自动驾驶"
        "传感器市场在2024年已达到约382亿美元的规模，预计到2030年将超过3200亿美元，2025-2030年复合年增长率"
        "（CAGR）约为30%-60%。其中，中国作为全球最大的智能网联汽车单一市场，表现尤为突出。"
    )

    # Market size table
    add_table(doc,
        headers=["年份", "市场规模（亿元）", "激光雷达占比", "毫米波雷达占比", "摄像头占比"],
        rows=[
            ["2023", "约260", "22%", "38%", "40%"],
            ["2025", "约380", "28%", "35%", "37%"],
            ["2027（预测）", "约490", "36%", "31%", "33%"],
            ["2030（预测）", "约980-1000", "45%", "25%", "30%"],
        ],
        caption="表1-1  中国自动驾驶传感器市场规模及预测",
        font_size=9
    )

    # Insert market size figure
    insert_image(doc, "market_size.png", width_inches=5.0,
                 caption="图1-1  中国自动驾驶传感器市场规模及预测（2023-2030）")

    # Insert market share figure
    insert_image(doc, "market_share.png", width_inches=5.0,
                 caption="图1-2  中国自动驾驶传感器市场份额分布（2025 vs 2030E）")

    key_insight_box(doc,
        "2025-2030年中国自动驾驶传感器市场CAGR约21%，激光雷达占比将从28%跃升至45%，"
        "成为市场份额最大的单一传感器品类。到2030年，中国预计将占全球自动驾驶传感器市场近40%的份额。"
    )

    heading_section(doc, "1.3  核心传感器细分市场分析")

    heading_subsection(doc, "1.3.1  激光雷达市场")
    body_paragraph(doc,
        "2024年全球激光雷达市场规模已达约512亿元（人民币），预计2027年将突破千亿元大关。中国厂商"
        "禾赛科技、速腾聚创已在全球车载前装市场中占据主导地位，2024年国产激光雷达全球市场占有率约达"
        "45%-50%。单车价值量方面，主激光雷达模组成本已从2020年的数万元降至2024年的3000-5000元区间，"
        "预计2027年将进一步下探至1500元以内。"
    )

    heading_subsection(doc, "1.3.2  毫米波雷达市场")
    body_paragraph(doc,
        "2024年中国车载毫米波雷达前装搭载量已超过1800万颗，同比增长27.6%，市场规模约86亿元。4D成像"
        "毫米波雷达自2024年起进入大规模商用元年，以增加的俯仰角分辨率弥补了传统3D毫米波雷达对静止物体"
        "的检测短板。77GHz CMOS毫米波雷达芯片已初步实现国产化替代，加特兰微电子（Calterah）的Alps系列"
        "SoC已实现大规模量产。"
    )

    heading_subsection(doc, "1.3.3  车载摄像头市场")
    body_paragraph(doc,
        "车载摄像头是单车搭载数量最多的传感器。L2级车辆平均搭载5-8颗摄像头，L3-L4级则可达到12-15颗"
        "以上。2024年全球车载摄像头市场规模约170亿美元，CMOS图像传感器（CIS）占据核心成本主导。中国厂商"
        "韦尔股份（OmniVision）、格科微等在车载CIS领域已具备全球竞争力，在800万像素高分辨率车载CIS产品"
        "线上持续推进国产替代。"
    )

    heading_section(doc, "1.4  典型企业与岗位分析")

    heading_subsection(doc, "1.4.1  其域创新（XGRIDS）")
    body_paragraph(doc,
        "其域创新（XGRIDS, https://xgrids.cn）是一家专注于三维空间数字化技术的公司，以自研的固态"
        "激光雷达和多传感器融合算法为核心竞争力。其产品线覆盖高精度手持三维扫描仪、机载LiDAR系统及"
        "自动驾驶感知解决方案。该公司在激光雷达3D重建与多传感器融合领域有深厚的工程积累，其招聘岗位"
        "直接反映了产业界对传感器技术复合型人才的具体需求。"
    )

    body_paragraph(doc,
        "（1）多传感器标定算法工程师：负责激光雷达、摄像头、IMU、GNSS等多传感器之间的内参和外参标定"
        "算法研发。核心工作包括基于PnP/手眼标定的LiDAR-Camera联合标定，多激光雷达之间的点云配准与"
        "SLAM轨迹优化。要求精通非线性优化理论（Ceres/G2O）、多视图几何及C++/Python编程。"
        "（岗位链接：https://pecivkvtit.jobs.feishu.cn/index/position/7486766416390113587/detail）"
    )

    body_paragraph(doc,
        "（2）感知融合算法工程师：负责多模态感知数据的融合与目标检测算法开发。核心工作包括基于"
        "BEV/Top-down视角的多传感器特征级融合，多目标跟踪（MOT）与轨迹预测算法，深度学习模型在"
        "嵌入式平台上的轻量化部署。要求熟悉PointPillars/CenterPoint等点云检测框架及PyTorch/TensorRT"
        "生态。（岗位链接：https://pecivkvtit.jobs.feishu.cn/index/position/7491253117397387529/detail）"
    )

    body_paragraph(doc,
        "（3）多传感器融合算法工程师（视觉雷达联合重建方向）：聚焦于视觉与LiDAR的联合三维场景重建，"
        "包括NeRF/3D Gaussian Splatting等前沿方法，用于高精地图构建与数字孪生。该岗位代表了传感器"
        "感知与三维视觉交叉的最新前沿方向。（岗位链接：https://pecivkvtit.jobs.feishu.cn/index/"
        "position/7535787737665538350/detail）"
    )

    heading_subsection(doc, "1.4.2  禾赛科技与速腾聚创")
    body_paragraph(doc,
        "禾赛科技（Hesai Technology）和速腾聚创（RoboSense）是全球车载激光雷达出货量排名前两位的企业。"
        "禾赛AT128（128线混合固态激光雷达）2023年单年出货超40万台，创造全球纪录。速腾聚创的E1平台采用"
        "自研SPAD阵列与VCSEL发射方案，实现了低于500美元的系统成本。这些企业的人才需求涵盖光电芯片设计、"
        "ASIC电路设计、点云算法、嵌入式系统工程、自动驾驶系统集成等方向，为光电信息科学与工程、电子科学"
        "与技术、自动化、计算机科学等专业背景的毕业生提供了广阔的职业通道。"
    )

    add_page_break(doc)

    # =====================================================================
    # 5. CHAPTER 2: 视觉摄像机传感器
    # =====================================================================
    heading_chapter(doc, "第二章  视觉摄像机传感器")

    heading_section(doc, "2.1  工作原理")

    heading_subsection(doc, "2.1.1  光电转换的物理基础")
    body_paragraph(doc,
        "车载视觉摄像机是自动驾驶感知系统中最基础、最成熟的传感器。其核心是一块CMOS（互补金属氧化物"
        "半导体）图像传感器芯片，基于内光电效应（Internal Photoelectric Effect）实现光信号到电信号的"
        "转换。当光子入射到半导体材料（通常为硅）表面时，若光子能量E = hν大于硅的禁带宽度Eg ≈ 1.12 eV，"
        "则价带中的电子被激发跃迁至导带，形成电子-空穴对。在PN结耗尽层中，内建电场将光生电子-空穴对分离，"
        "产生光电流Iph。光电流与入射光功率Popt的关系为：Iph = R × Popt = (ηq/hν) × Popt，式中R为响应度"
        "（A/W），η为量子效率，q为电子电荷量。"
    )
    add_formula_img(doc, "Iph", width_inches=4.2)

    heading_section(doc, "2.2  核心元器件")

    heading_subsection(doc, "2.2.1  CMOS图像传感器芯片架构")
    body_paragraph(doc,
        "车载CMOS图像传感器（CIS）芯片是一个高度集成的混合信号系统，其架构包含以下几个核心功能层："
        "像素阵列（Pixel Array）位于芯片中心，由数百万个独立像素以行-列矩阵排列构成，每个像素包含"
        "光电二极管和像素内读出晶体管。行驱动器（Row Driver）逐行选通像素，控制曝光和读出时序。列级"
        "模拟前端（Column Analog Front-End）对每一列的像素输出进行信号调理、CDS降噪和模数转换。"
        "数字逻辑模块（Digital Logic）包括时序控制器、图像信号处理器（ISP）、MIPI D-PHY高速串行接口"
        "等，负责图像数据的格式化和输出。"
    )

    heading_subsection(doc, "2.2.2  Pinned光电二极管与微透镜阵列")
    body_paragraph(doc,
        "现代车载CMOS图像传感器普遍采用4T-APS（4-Transistor Active Pixel Sensor）结构。每个像素包含"
        "四个晶体管：复位管（RST）、传输管（TX）、源跟随器（SF）和行选择管（RS），以及一个pinned光电"
        "二极管（PPD）。PPD的引入有效消除了传统PN结光电二极管的复位噪声（kTC噪声），使相关双采样（CDS）"
        "技术成为可能。3T-APS结构仅包含RST、SF和RS三个晶体管，没有传输管TX，无法实现CDS降噪，因此"
        "噪声性能劣于4T-APS，但在某些对成本极度敏感的应用中仍有使用。"
    )
    body_paragraph(doc,
        "除像素晶体管外，CMOS图像传感器的关键元器件还包括：微透镜阵列（Micro-Lens Array, MLA），每个"
        "像素上方集成一个微透镜，将入射光汇聚到光电二极管的光敏区域，提高填充因子和量子效率；彩色滤光片"
        "阵列（Color Filter Array, CFA），采用Bayer模式（RGGB）或其他定制模式，使不同像素对不同波段的"
        "可见光选择性响应，经后续去马赛克（Demosaic）算法重建彩色图像；红外截止滤光片（IR-Cut Filter），"
        "放置于传感器封装窗口处，滤除>650nm的近红外光以保证色彩还原的准确性。"
    )

    heading_section(doc, "2.3  测量电路")

    heading_subsection(doc, "2.3.1  相关双采样（CDS）降噪电路")
    body_paragraph(doc,
        "CMOS图像传感器的主要噪声源包括像素复位噪声（kTC噪声）、固定模式噪声（FPN）和暗电流散粒噪声。"
        "CDS技术通过对同一像素在复位电平（参考信号）与信号电平之间进行两次采样并做差，能够有效消除"
        "复位噪声和像素间的固定模式噪声：Vout = Vsignal - Vreset。这一差分操作是CMOS图像传感器信号链中"
        "最核心的噪声抑制电路设计。其工作流程为：(1) 首先复位光电二极管并采样复位电平Vreset；(2) 然后"
        "曝光积分，光生电荷在PPD中累积；(3) 采样信号电平Vsignal；(4) 做差得到净信号电压。由于复位噪声"
        "在微秒级时间间隔内具有高度相关性，差分操作几乎完全消除该噪声分量。"
    )
    add_formula_img(doc, "CDS", width_inches=4.0)
    add_circuit_img(doc, "cds_circuit", width_inches=5.5)

    heading_subsection(doc, "2.3.2  列级ADC架构")
    body_paragraph(doc,
        "为了将模拟像素信号转换为数字图像数据，CMOS图像传感器在每个像素列上集成一个列级ADC（Column-"
        "Parallel ADC）。主流架构包括单斜ADC（Single-Slope ADC）和SAR ADC（逐次逼近型ADC）。单斜ADC"
        "结构简单、面积小，适用于高分辨率传感器，但转换速率受限于比较器时钟频率，典型的12-bit转换需要"
        "约10μs。SAR ADC功耗效率高、转换速度快，适用于高帧率应用，已广泛用于最新一代车载CIS产品，"
        "如OmniVision的OX08B40（8MP车载CIS）即采用了高速列级SAR ADC架构。"
    )

    heading_section(doc, "2.4  应用场景")
    body_paragraph(doc,
        "车载摄像头根据不同安装位置和功能需求，分为多个类别，各具特定的视场角、分辨率和应用场景。"
    )

    add_table(doc,
        headers=["类型", "视场角", "分辨率", "主要功能"],
        rows=[
            ["前视单目", "30°-60°", "2-8MP", "FCW、LDW、TSR、ACC"],
            ["前视双目/三目", "30°-120°", "2-8MP", "立体视觉测距、AEB"],
            ["环视（鱼眼）", "180°-190°", "1-3MP", "360°全景泊车、BSD"],
            ["DMS舱内摄像头", "60°-90°", "1-2MP", "驾驶员疲劳监测、分心检测"],
            ["后视摄像头", "130°-150°", "1-2MP", "倒车影像、RVC"],
        ],
        caption="表2-1  车载摄像头分类与典型应用",
        font_size=9
    )

    body_paragraph(doc,
        "从应用角度看，车载视觉摄像机面临三大核心技术挑战："
    )
    body_paragraph(doc,
        "第一，高动态范围（HDR）。车载场景需要同时处理夜间弱光（<1 lux）和日间强光（>100,000 lux），"
        "要求HDR达到120-140dB。主流技术方案包括多重曝光合成（SME-HDR）和分像素双转换增益（DCG-HDR）。"
        "SME-HDR通过在不同曝光时间下连续拍摄多帧图像后合成，但存在运动伪影问题；DCG-HDR通过在像素内"
        "集成高低两个转换增益路径同步读出，避免了运动伪影，但牺牲了部分填充因子。"
    )
    body_paragraph(doc,
        "第二，LED闪烁抑制（LFM）。LED交通信号灯和车灯以PWM方式工作（典型频率90-300Hz），若曝光时间"
        "不匹配，会导致图像中出现闪烁或黑块。车载CIS需要支持LFM模式，通过延长像素曝光时间覆盖多个LED"
        "周期从而避免捕捉到LED的关断状态。然而，长曝光会牺牲HDR性能，因此需要在HDR与LFM之间进行系统级"
        "权衡。"
    )
    body_paragraph(doc,
        "第三，温度稳定性。车载环境温度范围宽（-40°C至+105°C），暗电流随温度呈指数增长（温度每升高约"
        "8°C暗电流约翻一番）。在+105°C高温下，暗电流可能比室温条件下高2-3个数量级，导致图像信噪比严重"
        "恶化。车载CIS需要片上温度传感器和自适应暗电流补偿电路来维持全温度范围内的图像质量。"
    )

    add_page_break(doc)

    # =====================================================================
    # 6. CHAPTER 3: 毫米波雷达传感器
    # =====================================================================
    heading_chapter(doc, "第三章  毫米波雷达传感器")

    heading_section(doc, "3.1  工作原理")

    heading_subsection(doc, "3.1.1  FMCW信号模型")
    body_paragraph(doc,
        "车载毫米波雷达是自动驾驶感知系统中不可或缺的全天候传感器。其工作频段主要位于24GHz（窄带，短距）"
        "和76-81GHz（宽带，长距）两个频段。国际上，24GHz频段正逐步被替代，77GHz和79GHz已成为车载毫米波"
        "雷达的主流频段。主流车载毫米波雷达采用调频连续波（Frequency Modulated Continuous Wave, FMCW）"
        "体制。发射信号为线性调频信号（chirp）：sTX(t) = AT cos[2π(fCt + (B/(2Tc))t²) + φ₀]，其中fC为"
        "起始频率（77GHz），B为调频带宽（通常为200MHz-4GHz），Tc为chirp周期。"
    )

    heading_subsection(doc, "3.1.2  混频与中频信号处理")
    body_paragraph(doc,
        "发射信号经目标反射后，接收信号相对于发射信号存在时延τ = 2R/c。接收信号与发射信号的本振（LO）"
        "副本在混频器（Mixer）中相乘，经低通滤波后得到中频（IF）信号。中频频率fIF包含目标的距离信息："
        "fIF = (2B × R) / (c × Tc)。由此可反演目标距离：R = (c × Tc × fIF) / (2B)。当目标具有径向相对"
        "速度vr时，根据多普勒效应，回波载频偏移fd = 2vr/λ。通过发射多个chirp形成「快时间-慢时间」二维信号"
        "矩阵，经二维FFT（Range-Doppler FFT）即可同时提取目标的距离和速度信息。第一维FFT（快时间维FFT）"
        "分离不同距离的目标，每个距离bin对应一个特定fIF；第二维FFT（慢时间维FFT）在每个距离bin内分离不同"
        "速度的目标。"
    )

    heading_subsection(doc, "3.1.3  角度估计与4D成像雷达")
    body_paragraph(doc,
        "传统毫米波雷达通过多个接收天线组成的虚拟阵列，利用到达角（AoA）估计目标的方位角。对于包含NTX个"
        "发射天线和NRX个接收天线的MIMO阵列，可形成等效NTX × NRX个虚拟接收通道。到达角θ通过相邻天线间的"
        "相位差估计：θ = arcsin[(λ × Δφ)/(2π × d)]，其中d为天线间距，Δφ为相邻通道间的相位差。"
    )
    body_paragraph(doc,
        "传统的3D毫米波雷达仅能获得目标的距离-方位角-多普勒信息。而4D成像雷达通过增加俯仰向天线阵列"
        "（MIMO垂直维度），额外获得目标的高度信息，使雷达拥有「成像级」的点云输出能力。4D成像雷达可输出"
        "数万点/帧的稀疏点云，有效识别静止障碍物、路沿和桥梁，弥补了传统3D毫米波雷达对静止物体检测能力"
        "不足的核心短板。2024-2025年，大陆集团ARS540、华为4D成像雷达、德赛西威等产品已实现大规模前装量产。"
    )

    heading_section(doc, "3.2  核心元器件")

    heading_subsection(doc, "3.2.1  77GHz MMIC收发芯片")
    body_paragraph(doc,
        "毫米波雷达的核心是单片微波集成电路（MMIC），集成了以下关键功能模块："
    )
    body_paragraph(doc,
        "频率合成器/锁相环（PLL）利用分数N分频锁相环产生高线性度的线性调频信号。chirp的线性度直接决定了"
        "雷达的距离分辨率ΔR = c/(2B)。chirp非线性会导致中频信号展宽，降低距离分辨率和信噪比。因此，PLL的"
        "相位噪声和线性度是MMIC设计的核心指标。功率放大器（PA）将微弱的chirp信号放大至足够功率后馈送至"
        "发射天线，典型发射功率约12-15dBm。"
    )
    body_paragraph(doc,
        "低噪声放大器（LNA）位于接收链路第一级，其噪声系数（NF）决定了雷达的整体灵敏度。车载77GHz LNA的"
        "典型噪声系数为4-6dB。混频器（Mixer）将毫米波信号下变频至基带/中频，其转换增益和线性度影响IF信号"
        "的保真度。混频后的IF信号经可编程增益放大器（PGA）和抗混叠低通滤波器后，由ADC进行数字化采样，"
        "典型ADC分辨率为12-16bit，采样率20-50MSPS。"
    )

    heading_subsection(doc, "3.2.2  MIMO天线阵列与CMOS国产替代")
    body_paragraph(doc,
        "近年来，77GHz CMOS毫米波雷达芯片取得了重大突破。传统的SiGe BiCMOS工艺虽然射频性能优异（fT/fmax"
        "可达300/500GHz），但成本高且难以与数字基带集成。CMOS工艺（如28nm/45nm RF-CMOS）使MMIC与数字信号"
        "处理（DSP/MCU）的单芯片集成（SoC）成为可能，大幅降低了雷达模组的成本和小型化程度。"
    )
    body_paragraph(doc,
        "国产厂商加特兰微电子（Calterah）的77GHz/79GHz Alps系列CMOS毫米波雷达SoC已实现大规模量产供货，"
        "集成了3发4收MMIC通道、DSP和MCU于单一芯片。森思泰克、楚航科技等模组厂商也相继推出了基于国产芯片"
        "方案的4D成像雷达产品，推动车载毫米波雷达国产化率从2021年的不足10%提升至2025年的约35%，"
        "预计2030年将突破80%。"
    )

    heading_section(doc, "3.3  测量电路")

    heading_subsection(doc, "3.3.1  PLL频率合成器与Chirp生成")
    body_paragraph(doc,
        "FMCW雷达的测距精度和分辨率高度依赖于chirp信号的线性度。频率合成器的核心是一个分数N分频"
        "锁相环（Fractional-N PLL），通过Σ-Δ调制器控制多模分频器的分频比，产生精确的线性调频斜坡。"
        "chirp线性度通常用频率偏差的均方根（RMS Frequency Error）来衡量，典型值需控制在<0.1%以内。"
        "非线性会导致中频信号频谱展宽，降低距离分辨率并引入虚假目标。PLL环路滤波器的带宽设计需要在"
        "相位噪声抑制和chirp跟踪速度之间取得平衡——带宽过低会导致chirp起始段的频率过冲，带宽过高则"
        "无法充分抑制VCO的相位噪声。现代77GHz PLL通常采用两点调制（Two-Point Modulation）架构，"
        "在VCO调谐端口和PLL反馈路径同时注入调制信号，实现宽带低噪声的chirp生成。"
    )

    heading_subsection(doc, "3.3.2  混频器与中频链路")
    body_paragraph(doc,
        "接收链路中，LNA放大后的微弱回波信号（典型-100至-60dBm）与PA耦合出的本振信号在混频器中"
        "下变频至基带中频。混频器的转换增益（Conversion Gain, 典型10-15dB）和线性度（IIP3> 0dBm）"
        "直接影响IF信号的信噪比。混频后的IF信号首先经高通滤波器滤除由本振泄漏和自混频产生的直流偏移，"
        "然后经可编程增益放大器（PGA, 增益范围20-60dB）进行动态范围压缩——近距离强反射目标的回波功率"
        "可能比远距离弱目标高80dB以上。抗混叠低通滤波器（截止频率通常为10-25MHz，对应最大探测距离"
        "150-250m）滤除高频噪声后，由12-16bit ADC以20-50MSPS采样率转换为数字信号，进入后续的"
        "Range-Doppler二维FFT处理器。"
    )

    add_circuit_img(doc, "fmcw_trx_circuit", width_inches=5.5)

    heading_section(doc, "3.4  应用场景")
    body_paragraph(doc,
        "毫米波雷达的核心优势包括：全天候工作能力强（毫米波在雨、雾、雪、烟尘中的衰减远小于激光和可见光）；"
        "速度测量直接且精准（通过多普勒效应直接获取目标径向速度，无需帧间差分计算）；成本相对较低（单颗"
        "77GHz前雷达成本已降至200-400元区间）。主要局限包括：对静止物体的检测能力弱（4D成像雷达正在改善）；"
        "角度分辨率有限（受天线孔径限制，方位角分辨率通常为1°-3°，远低于激光雷达的0.1°级）；对非金属目标"
        "（行人等）反射较弱。"
    )

    add_page_break(doc)

    # =====================================================================
    # 7. CHAPTER 4: 超声波雷达传感器
    # =====================================================================
    heading_chapter(doc, "第四章  超声波雷达传感器")

    heading_section(doc, "4.1  工作原理")

    heading_subsection(doc, "4.1.1  压电效应与换能机制")
    body_paragraph(doc,
        "超声波雷达（Ultrasonic Sensor）是汽车电子系统中应用最广泛的近距离传感器。其核心工作原理基于压电"
        "效应（Piezoelectric Effect）与脉冲回波法（Pulse-Echo Method）。超声波传感器的核心换能元件为压电"
        "陶瓷片（如PZT-4/PZT-8系列锆钛酸铅陶瓷）。当在压电陶瓷两侧施加交变电压时，陶瓷片将产生与电压频率"
        "一致的机械振动，从而向空气中辐射超声波（逆压电效应）。当接收到的超声回波作用于陶瓷片上时，机械"
        "振动转换为电荷输出（正压电效应）。"
    )
    body_paragraph(doc,
        "车载超声波传感器通常工作于40-58kHz频段。该频段的选择是探测距离、方向性和抗环境噪声之间的折中。"
        "频率越高，方向性越好（波束角越窄），但空气中衰减越大；频率越低，探测距离越远，但方向性差且易受"
        "环境噪声干扰。40kHz是使用最广泛的频率，在空气中波长约8.5mm，可较好地兼顾距离与分辨率。"
    )

    heading_subsection(doc, "4.1.2  渡越时间测距模型")
    body_paragraph(doc,
        "传感器发射一组超声波脉冲（通常为8-16个周期），脉冲经空气传播至障碍物后反射回来。设发射时刻为t0，"
        "回波到达时刻为t1，声速为vs（空气中约340m/s），则障碍物距离为：R = 0.5 × vs × (t1 - t0) = 0.5 ×"
        "vs × Δt。渡越时间Δt的典型范围约为0.17-41ms（对应距离约0.03-7m）。为提高测量可靠性，系统通常发射"
        "多组脉冲并取回波时间的统计平均值。"
    )
    body_paragraph(doc,
        "声速vs是温度的函数：vs(T) = 331.3 × √(1 + T/273.15) ≈ 331.3 + 0.606 × T (m/s)，其中T为环境温度"
        "（°C）。温度从-40°C变化至+85°C时，声速从约307m/s变化至约383m/s，相对变化约25%。若不补偿，将引入"
        "同等比例的距离误差。因此，车载超声波系统通常集成温度传感器（如NTC热敏电阻或片上温度传感器）以实时"
        "修正声速，将温度引入的距离误差控制在±1%以内。"
    )

    heading_section(doc, "4.2  核心元器件")

    heading_subsection(doc, "4.2.1  压电陶瓷换能器")
    body_paragraph(doc,
        "车载超声波传感器的核心换能元件为压电陶瓷片（PZT-4/PZT-8系列锆钛酸铅陶瓷）。PZT-4具有高机械"
        "品质因数（Qm>500）和低介电损耗，适合大功率发射；PZT-8的介电常数较高，灵敏度更优，适合收发型"
        "应用。换能器的谐振频率由陶瓷片的厚度和径向尺寸决定，40kHz换能器的典型尺寸为直径10-14mm、厚度"
        "约2-3mm。压电陶瓷片被封装于铝制或塑料壳体中，前端匹配层（λ/4厚度）用于声阻抗匹配（从PZT的"
        "~30MRayl过渡到空气的~430Rayl），以提高声波辐射效率。"
    )

    heading_subsection(doc, "4.2.2  ASIC集成驱动芯片")
    body_paragraph(doc,
        "现代APA（自动泊车辅助）系统中，超声波传感器趋于采用专用ASIC集成驱动芯片，如Elmos E524系列、"
        "TDK ICS-43434等。这类ASIC内部集成了升压DC-DC转换器（产生100-300V发射脉冲）、可编程脉冲发生器"
        "（控制发射脉冲数量和频率）、低噪声回波接收放大链路（包括LNA、BPF、VGA和阈值比较器）、12bit "
        "SAR ADC以及SPI/UART数字通信接口。ASIC方案将传统需要10-15颗分立元件的传感器驱动电路压缩至"
        "单一芯片（典型封装QFN-20, 4×4mm），大幅减少了PCB面积，为传感器的小型化和一致性提供了保证。"
    )

    heading_section(doc, "4.3  测量电路")

    heading_subsection(doc, "4.3.1  发射驱动电路")
    body_paragraph(doc,
        "超声波发射需要数百伏的瞬时高压脉冲以驱动压电陶瓷产生足够的声压级（典型>100dB SPL@30cm）。典型"
        "电路采用变压器升压拓扑：利用中心抽头变压器的匝数比（通常1:10至1:50），将MCU输出的低压脉冲（3.3V/"
        "5V）升压至100-300Vpp驱动换能器。变压器方案结构简单可靠，但体积较大。"
    )
    body_paragraph(doc,
        "现代自动泊车系统（APA）中的超声波传感器趋于采用ASIC集成驱动芯片，如Elmos E524系列，内部集成了"
        "升压DCDC转换器、脉冲发生器、回波接收放大链路、ADC和数字信号处理单元。ASIC方案大幅减少了外围元件"
        "数量，单个传感器模块的PCB面积可压缩至20×15mm以内，满足了现代汽车对传感器小型化的严格要求。"
    )

    heading_subsection(doc, "4.3.2  回波接收与阈值检测")
    body_paragraph(doc,
        "接收链路包括四个关键模块。低噪声前置放大器（LNA）将压电换能器输出的微伏至毫伏级回波信号进行初级"
        "放大，增益通常为20-40dB，噪声系数需小于5nV/√Hz。带通滤波器（BPF）中心频率40-58kHz，带宽2-5kHz，"
        "用于抑制发动机噪声（<1kHz）、风噪（宽带）和电磁干扰等带外干扰。可变增益放大器（VGA）补偿信号随距离"
        "增加而指数衰减的物理特性——声波在空气中按球面波扩展，声压随距离平方反比衰减，加上空气吸收损耗，"
        "总衰减大约为1/R² × e^(-αR)。阈值比较器将处理后的模拟信号与设定阈值比较，阈值通常随时间动态变化"
        "（时间-增益控制TGC），近距时阈值较高以避免振铃误触发，远距时阈值逐渐降低以提高灵敏度。"
    )

    add_circuit_img(doc, "ultrasonic_circuit", width_inches=5.5)

    heading_section(doc, "4.4  应用场景")
    body_paragraph(doc,
        "超声波雷达广泛应用于三类场景：UPA（泊车辅助），探测距离0.3-2.5m，传感器安装于前后保险杠；APA"
        "（自动泊车辅助），探测距离0.3-5m，需要更多传感器（8-12颗）和侧向安装；BSD（盲区检测辅助），使用"
        "长距超声波传感器，探测距离可达5-7m。"
    )
    body_paragraph(doc,
        "超声波雷达的性能边界包括：探测距离短（≤7m），受限于空气中超声波衰减；方向性差（波束角约60°-120°），"
        "难以精确定位障碍物的空间方位；受风噪/雨滴撞击干扰大，恶劣天气下虚警率升高；声速因温度变化引入距离"
        "误差；多个传感器之间可能产生串扰（Crosstalk），需要时分复用或编码调制来区分不同传感器的回波。"
        "然而，凭借其极低的成本（单颗传感器成本约5-15元人民币），超声波雷达在低速近距离探测场景中仍是"
        "不可替代的传感器方案。"
    )

    add_page_break(doc)

    # =====================================================================
    # 8. CHAPTER 5: 激光雷达传感器
    # =====================================================================
    heading_chapter(doc, "第五章  激光雷达（LiDAR）传感器")

    body_paragraph(doc,
        "激光雷达（Light Detection and Ranging, LiDAR）是L3+自动驾驶感知系统的核心传感器，也是当前技术"
        "迭代最快、资本密集度最高的传感器赛道。本章从物理原理到测量电路进行系统性深入分析。"
    )

    heading_section(doc, "5.1  工作原理")

    heading_subsection(doc, "5.1.1  内光电效应与单光子探测")
    body_paragraph(doc,
        "激光雷达接收端的物理基础是内光电效应。当波长为λ的激光光子入射到半导体探测器表面时，若光子能量"
        "E = hc/λ大于半导体材料的禁带宽度Eg，价带电子被激发至导带，形成电子-空穴对。在外部反向偏置高压的"
        "作用下，这些光生载流子在耗尽层强电场中被加速漂移，产生可测量的光电流。对于1550nm波长的激光，光子"
        "能量E ≈ 0.80 eV，InGaAs材料的Eg ≈ 0.75 eV，可以满足内光电效应的能量条件。而对于905nm波长，"
        "E ≈ 1.37 eV，硅基材料（Eg ≈ 1.12 eV）完全适用。"
    )

    heading_subsection(doc, "5.1.2  飞行时间法（ToF）测距模型")
    body_paragraph(doc,
        "ToF（Time of Flight）是当前量产固态激光雷达最成熟的测距体制。发射端发出一束极窄的脉冲激光（脉宽"
        "τp通常为1-10ns），目标反射的回波经自由空间传播后到达接收端。距离R由激光往返时间Δt计算：R = 0.5 ×"
        "c × Δt，其中c = 3×10⁸ m/s为光速。要达到1cm的测距精度，Δt的测量精度需达到约66.7ps（皮秒）——这对"
        "时间测量电路提出了极高的要求。ToF体制的优势在于结构相对简单、峰值功率高、抗环境光能力强，是当前"
        "905nm和1550nm车载LiDAR的主流方案。"
    )

    heading_subsection(doc, "5.1.3  FMCW激光雷达")
    body_paragraph(doc,
        "FMCW体制的激光雷达与FMCW毫米波雷达共享相同的数学模型基础，但在光学域实现。发射激光的频率随时间"
        "线性变化（通过调制激光器的注入电流或使用外部调制器），回波与本振光在光学混频器中产生拍频。目标距离"
        "和速度分别由拍频的中频分量和多普勒频移决定：fIF = (2B×R)/(c×Tm)，fd = 2vr/λ。FMCW激光雷达的核心"
        "优势在于：仅对与本振光频率匹配的信号产生响应，天然免疫其他激光雷达的脉冲串扰和阳光干扰；同时测距"
        "测速，无需帧间差分；相干接收的散粒噪声极限远低于直接探测。其技术瓶颈在于需要窄线宽（kHz级）、频率"
        "可线性调谐的激光器，以及高带宽的光学混频和平衡探测器，系统复杂度和成本显著高于ToF方案。"
    )

    heading_section(doc, "5.2  核心元器件")

    heading_subsection(doc, "5.2.1  发射端：从EEL到多结VCSEL")
    body_paragraph(doc,
        "边发射激光器（EEL）是传统方案，光束从芯片侧面出射，难以形成二维阵列，适合旋转式/转镜式机械扫描"
        "雷达。垂直腔面发射激光器（VCSEL）光束垂直于芯片表面出射，可实现大规模二维阵列的晶圆级制造和测试，"
        "是固态Flash激光雷达和可寻址VCSEL阵列的核心光源。905nm VCSEL（GaAs基）已实现多结（Multi-Junction）"
        "堆叠——通过在同一外延结构中生长多个有源区并串联连接，在保持人眼安全等级的前提下将峰值功率密度提升"
        "3-5倍，使固态Flash激光雷达的探测距离从数十米提升至200m以上。"
    )
    body_paragraph(doc,
        "波长选择存在905nm与1550nm两大技术路线的竞争。1550nm（InP基）的光源远离人眼视网膜吸收带，允许发射"
        "更高的激光功率，因此探测距离更远。但其探测器（InGaAs APD/SPAD）成本高于硅基方案。905nm方案凭借"
        "硅基探测器的成本优势在现阶段的中短距ADAS市场占据主导，而1550nm方案在L4 Robotaxi等需要300m+长距"
        "探测的场景中具有不可替代的优势。"
    )

    heading_subsection(doc, "5.2.2  接收端：从PIN到SPAD/SiPM")
    body_paragraph(doc,
        "探测器从普通PIN光电二极管到雪崩光电二极管（APD），再到单光子雪崩二极管（SPAD），其灵敏度逐步提高"
        "了数个数量级。PIN光电二极管无内部增益，响应度0.4-0.8 A/W（Si@905nm），仅适用于近距离强回波场景。"
        "APD工作于线性模式，施加的反向偏压低于击穿电压，光生载流子在耗尽层高电场中被加速，通过碰撞电离产生"
        "链式倍增（雪崩效应），增益M通常为50-200倍，噪声以过剩噪声因子F(M) = M^x（x≈0.3-0.5 for Si）表征。"
    )
    body_paragraph(doc,
        "SPAD工作于Geiger模式，偏压高于击穿电压，此时APD进入非线性的雪崩击穿状态——单个入射光子即可触发"
        "出宏观的雪崩电流（mA级），探测灵敏度达到量子极限。SiPM（硅光电倍增管）由成千上万个独立SPAD微单元"
        "并联组成，每个微单元串联一淬灭电阻。SiPM的宏观输出电流正比于同时触发雪崩的微单元数量，从而在保留"
        "SPAD单光子灵敏度的同时获得了光子数分辨能力（动态范围可达10³-10⁴个光子/脉冲）。"
    )

    heading_subsection(doc, "5.2.3  固态扫描器件：MEMS与OPA")
    body_paragraph(doc,
        "实现三维空间覆盖需要光束扫描机制。MEMS微振镜利用MEMS工艺制造的微米级硅基反射镜，在电磁或静电力"
        "驱动下发生二维偏转，改变光路方向。MEMS振镜已成功应用于禾赛AT128（一维转镜+MEMS二维扫描）等量产"
        "产品，反射镜的谐振频率通常为1-2kHz（快轴）和几十Hz（慢轴），可实现类似利萨如（Lissajous）或光栅"
        "扫描图案。光学相控阵（OPA）利用集成光波导阵列中的热光或电光效应，通过改变各波导通道中光的相对"
        "相位，在远场利用多光束干涉原理实现光斑的无惯性扫描，是理想的「纯固态」方案。但当前OPA受限于波导"
        "损耗、串扰和光栅瓣等物理问题，尚不能满足车载全视场扫图的需求。"
    )

    heading_section(doc, "5.3  测量电路")

    heading_subsection(doc, "5.3.1  跨阻放大器（TIA）")
    body_paragraph(doc,
        "光电探测器（APD/SPAD/SiPM）输出的信号是极其微弱的电流脉冲（典型峰值在微安至毫安量级，脉宽在纳秒"
        "级）。跨阻放大器（Transimpedance Amplifier, TIA）是连接探测器与后续数字电路之间的关键桥梁，其基本"
        "功能是将微弱电流信号转换为可被ADC或TDC处理的电压信号。理想TIA的输出为：Vout = -Iin × Rf，其中Rf为"
        "反馈电阻。"
    )
    body_paragraph(doc,
        "在实际高频电路中，光电二极管的结电容Cd与反馈电阻Rf形成RC低通极点，限制了TIA的带宽。要保证纳秒级"
        "脉冲的保真度，TIA的闭环带宽f-3dB需满足：f-3dB ≥ 0.35/tr（其中tr为脉冲上升时间）。对于tr=1ns的"
        "脉冲，TIA带宽需达到350MHz以上。为实现高增益与高带宽的兼顾，需要在Rf两端并联补偿电容Cf以引入零点"
        "补偿。TIA的闭环传递函数为：H(s) = Vout(s)/Iin(s) = -Rf/(1 + s×Rf×Cf)。通过选择Cf使得Rf×Cf匹配"
        "输入端的RC时间常数，可获得最大平坦带宽（Butterworth响应）。"
    )

    add_table(doc,
        headers=["文献", "工艺节点", "带宽", "输入噪声", "功耗/通道"],
        rows=[
            ["方吉鑫等, 2026", "180nm CMOS", "900MHz", "2.5pA/√Hz", "1.5mW"],
            ["林远启等, 2021", "SiGe BiCMOS", "200MHz", "3.2pA/√Hz", "3.5mW"],
            ["某国际厂商", "65nm CMOS", "750MHz", "1.8pA/√Hz", "1.2mW"],
        ],
        caption="表5-1  先进TIA集成芯片性能对比",
        font_size=9
    )

    body_paragraph(doc,
        "TIA的等效输入噪声电流谱密度主要由三部分构成：反馈电阻热噪声（4kT/Rf）、运放输入电压噪声"
        "通过寄生电容的增益放大效应、以及运放输入电流噪声。增大Rf可降低反馈电阻热噪声贡献，但同时"
        "增加了输入电压噪声的耦合放大效应——这揭示了TIA设计的核心矛盾。在先进CMOS工艺中，采用共源"
        "共栅（Cascode）输入级和噪声抵消技术，可将等效输入噪声压低至1.5-3.0pA/√Hz量级。"
    )
    body_paragraph(doc,
        "以905nm激光雷达为典型场景进行信噪比链路预算：设发射峰值功率Ppeak=100W，目标距离R=100m，"
        "目标反射率ρ=0.1（朗伯反射体），接收光学孔径D=20mm。根据激光雷达方程计算得接收功率Prx约"
        "157nW。若SPAD的探测效率（PDE）为10%，则到达探测器的光子数率约为7.1×10^10 photons/s。"
        "TIA需在此动态范围内保持线性放大，同时保证输出SNR≥10以满足后续TDC的触发阈值判别需求。"
    )

    add_formula_img(doc, "TIA", width_inches=3.5)
    add_formula_img(doc, "TIA_noise", width_inches=5.5)
    add_circuit_img(doc, "tia_circuit", width_inches=5.2)

    heading_subsection(doc, "5.3.2  时间数字转换器（TDC）")
    body_paragraph(doc,
        "TDC（Time-to-Digital Converter）是ToF激光雷达测距精度的决定性电路。对于1cm的距离分辨率需求，时间"
        "测量精度必须达到约67ps。传统基于时钟计数器的方法受限于系统时钟频率（几百MHz量级，对应几纳秒分辨率），"
        "无法满足皮秒级精度要求。"
    )
    body_paragraph(doc,
        "抽头延迟线型TDC利用基本门电路（缓冲器/反相器）的传播延迟作为精细时间刻度。Start信号在一串串联"
        "缓冲器（延迟线）中传播，每个缓冲器产生固定的小延迟τLSB（在先进CMOS工艺中可达10-20ps）。当Stop信号"
        "（回波到达）到来时，触发一组D触发器锁存延迟线上每个抽头的当前逻辑状态，设被记录的「1」的个数为N，"
        "则精细时间间隔为：Δt_fine = N × τLSB。Vernier型TDC则采用两条具有微小延迟差的延迟线（如τ1=50ps，"
        "τ2=48ps），利用游标卡尺的「对齐」原理将分辨率进一步提高至τ1-τ2=2ps量级。"
    )

    body_paragraph(doc,
        "TDC的单次测量精度受多个误差源限制，总计时抖动σTDC可建模为量化误差、延迟线微分非线性"
        "（DNL）的累积偏差、工艺-电压-温度（PVT）波动引起的延迟漂移、以及回波脉冲幅度变化导致的"
        "定时游走（Time Walk）误差的均方根合成。由TDC计时抖动通过ToF方程传播至距离误差为"
        "σR = (c/2)×σTDC。以σTDC=30ps为例，σR≈4.5mm。若采用多脉冲累加（如发射100个脉冲取"
        "ToF中值），测距精度可在白噪声假设下提升为σR/√100≈0.45mm。实际系统中常采用基于直方图的"
        "统计算法（TCSPC），通过构建光子到达时间统计直方图，以峰值检测或加权质心法提取高置信度"
        "的ToF估计值。"
    )

    add_formula_img(doc, "TDC", width_inches=3.5)
    add_formula_img(doc, "TDC_jitter", width_inches=5.5)
    add_circuit_img(doc, "tdc_circuit", width_inches=5.5)

    heading_subsection(doc, "5.3.3  淬灭电路与暗电流抑制")
    body_paragraph(doc,
        "SPAD工作于Geiger模式，其对入射光子的响应具有概率统计特性。单个SPAD微单元在一次探测"
        "周期内发生雪崩的概率Pdet服从非齐次泊松过程：Pdet = 1 - exp(-∫[ηPDE×Φph(t) + DCR]dt)，"
        "式中ηPDE为光子探测效率（为量子效率ηQE与几何填充因子FF及雪崩触发概率Pav的乘积），"
        "Φph(t)为入射光子通量，DCR为暗计数率（单位：cps）。SiPM由Nmicro个独立SPAD微单元并联"
        "组成，其宏观输出电流正比于同时触发雪崩的微单元数。当触发数远小于总数时（线性区），"
        "输出正比于入射光子通量；当触发数接近总数时进入饱和区，输出-输入关系偏离线性，"
        "需在系统标定中予以补偿。"
    )

    add_formula_img(doc, "SPAD_prob", width_inches=5.5)
    add_formula_img(doc, "SiPM", width_inches=4.5)

    body_paragraph(doc,
        "SPAD在触发雪崩后需要主动淬灭和复位以准备下一次探测。无源淬灭（Passive Quenching）串联一个大电阻"
        "或MOSFET电流源，雪崩电流在电阻上的压降使偏压降至击穿电压以下，自然淬灭。结构简单，但恢复时间（Dead "
        "Time）较长（数十ns到数百ns），限制了最大光子计数率。主动淬灭（Active Quenching）通过高速反馈电路"
        "检测雪崩事件的发生，主动将偏压拉低，在雪崩结束后再主动恢复到工作偏压。恢复时间可压缩至几纳秒，"
        "适合高光子通量场景，但电路复杂度和功耗显著增加。"
    )

    body_paragraph(doc,
        "暗电流（Dark Current）是限制弱光探测器性能的核心物理瓶颈。即使在完全无光的条件下，反向偏置的PN结"
        "中仍存在因热激发（Shockley-Read-Hall产生-复合过程）和带间隧穿而产生的微弱电流。暗电流的主要来源为："
        "Idark ∝ ni ∝ T^(3/2) × exp(-Eg/(2kT))。暗电流对温度极为敏感——温度每升高约8-10°C，暗电流约翻一番。"
        "在车载高温环境（+85°C）下，暗电流可能比室温条件下高2-3个数量级，导致SPAD的暗计数率（DCR）急剧上升"
        "（从室温下的10-100cps升至高温下的10⁴-10⁵cps），产生大量虚假事件（噪声点云），严重恶化测距的置信度。"
        "噪声抑制策略包括：芯片制造环节减少缺陷和杂质陷阱态；系统层面通过TEC热电制冷器将探测器温度稳定在"
        "25°C附近；算法层面通过时间符合滤波剔除孤立暗计数触发。"
    )

    add_circuit_img(doc, "spad_quenching_circuit", width_inches=5.2)

    heading_subsection(doc, "5.3.4  SIP系统级封装与信号完整性")
    body_paragraph(doc,
        "为了抑制PCB长走线引入的寄生电感和寄生电容对微弱高速信号的衰减，现代激光雷达接收系统广泛采用系统级"
        "封装（System in Package, SIP）技术，将探测器芯片与TIA芯片在物理空间上进行微米级的超短互连，封装于"
        "同一管壳内。SIP不仅消除了中间连接线上的寄生参数（每毫米键合线约引入1nH寄生电感），而且减少了电磁"
        "干扰（EMI）接收面积。部分先进方案（如禾赛AT128的接收模块）将APD阵列、TIA阵列和初级ADC以3D堆叠"
        "（3D-Stacking）方式集成，采用TSV（硅通孔）实现垂直互连，将互连长度从毫米级缩短至微米级，极大提升"
        "了接收链路的信号完整性和通道间一致性。"
    )

    heading_section(doc, "5.4  应用场景")
    body_paragraph(doc,
        "激光雷达在自动驾驶中的应用覆盖L2+至L4多个级别。在L2+/L3级ADAS系统中，前向长距激光雷达"
        "（探测距离200-300m）作为摄像头和毫米波雷达的补充，提供高精度的三维环境感知，特别是在AEB"
        "（自动紧急制动）和HWP（高速公路领航）功能中显著降低了对静止异形障碍物的漏检率。在L4级"
        "Robotaxi中，多颗激光雷达（通常为顶部360°主雷达+四周补盲雷达，共4-6颗）构成360°无死角点云覆盖，"
        "是路径规划与避障决策的核心输入源。激光雷达还广泛应用于高精地图采集与更新、智慧城市V2X基础设施"
        "感知（路侧激光雷达）等领域。目前禾赛AT128、速腾聚创M1/E1等国产激光雷达已在理想、小鹏、蔚来、"
        "小米等品牌的多款量产车型上前装搭载，标志着车载激光雷达已从\"选装尝鲜\"进入\"前装标配\"的产业化阶段。"
    )

    add_page_break(doc)

    # =====================================================================
    # 9. CHAPTER 6: IMU惯性测量单元
    # =====================================================================
    heading_chapter(doc, "第六章  惯性测量单元（IMU）传感器")

    heading_section(doc, "6.1  工作原理")

    heading_subsection(doc, "6.1.1  MEMS加速度计")
    body_paragraph(doc,
        "惯性测量单元（Inertial Measurement Unit, IMU）通过感知载体的加速度和角速度来推算载体的运动状态，"
        "是自动驾驶定位系统中的关键传感器。现代车载IMU几乎全部基于微机电系统（MEMS）技术制造，以极小的"
        "体积（<1cm³）和极低的成本（<$10）提供战术级惯性测量能力。"
    )
    body_paragraph(doc,
        "MEMS加速度计的基本原理是将加速度引起的惯性力转换为电容或压阻变化。以最主流的电容式MEMS加速度计为"
        "例，其核心结构是一个通过弹性悬臂梁支撑的硅基质量块（Proof Mass），质量块两侧设有固定梳齿电极。当"
        "载体沿敏感轴方向产生加速度a时，惯性力F = m×a使质量块相对于固定电极产生位移Δx，引起差分电容C1与"
        "C2的变化：ΔC = C1 - C2 ∝ Δx ∝ a。电容变化通过片上开关电容（Switched-Capacitor）读出电路转换为"
        "电压信号，经Σ-Δ ADC后输出高分辨率数字加速度值。MEMS加速度计的关键参数包括测量范围（通常±2g至±16g）、"
        "分辨率（μg级）、带宽（>100Hz）和零偏稳定性。"
    )

    heading_subsection(doc, "6.1.2  MEMS陀螺仪与科里奥利力效应")
    body_paragraph(doc,
        "MEMS振动陀螺仪利用科里奥利力（Coriolis Force）效应测量角速度。其结构包含一个在驱动方向上以恒定"
        "频率ωd振动的质量块。当载体绕垂直于驱动-检测平面的轴以角速度Ω旋转时，质量块在检测方向上受到科里奥利"
        "力作用：Fc = -2m × (Ω × v)，其中m为质量块质量，v为质量块在驱动方向上的瞬时速度。科里奥利力使质量块"
        "在检测方向上产生与角速度Ω成正比的位移（通常为亚纳米至皮米级），由高灵敏度电容检测电路读出。为提高"
        "信噪比，驱动和检测回路均采用锁相放大（Phase-Locked Loop）技术，在谐振频率处实现高Q值机械放大。"
        "MEMS陀螺仪的典型角速度范围为±125°/s至±2000°/s，零偏不稳定性为1-10°/hr（工业级）或<1°/hr（战术级）。"
    )

    add_formula_img(doc, "Coriolis", width_inches=4.0)

    heading_section(doc, "6.2  核心元器件")

    heading_subsection(doc, "6.2.1  梳齿电容式MEMS加速度计")
    body_paragraph(doc,
        "电容式MEMS加速度计的核心结构是一个通过弹性悬臂梁支撑的硅基质量块（Proof Mass），质量块两侧设有"
        "固定梳齿电极。质量块尺寸通常为数百微米见方，厚度约10-50μm，由深反应离子刻蚀（DRIE）工艺在SOI"
        "（Silicon-On-Insulator）晶圆上加工而成。梳齿电极的指间距通常为1-3μm，单个加速度计包含数百对"
        "梳齿以增大敏感电容的绝对值和信噪比。当载体加速度使质量块产生位移时，差分电容C1与C2的变化量"
        "ΔC典型值为10-100fF/g，需由高精度电容读出电路检测。",
    )

    heading_subsection(doc, "6.2.2  MEMS振动陀螺仪")
    body_paragraph(doc,
        "MEMS振动陀螺仪包含驱动质量块和检测质量块两个机械子系统。驱动方向通过静电梳齿驱动器以锁相环"
        "（PLL）维持在谐振频率（通常为10-30kHz），实现稳定的恒幅振荡（振幅通常为几μm）。检测方向的质量块"
        "在科里奥利力作用下产生位移，由差分电容电极检测。为提高灵敏度，检测模态也被调谐至与驱动模态接近的"
        "频率（模态匹配），利用机械Q值放大（Q值典型为10,000-100,000 in vacuum）。MEMS陀螺仪的真空封装"
        "（<1mTorr）是保证高Q值的核心工艺——封装泄漏将导致Q值急剧下降和噪声增大。Bosch BMI270、InvenSense "
        "ICM-42688等6轴MEMS IMU将3轴加速度计和3轴陀螺仪集成于单一芯片（2.5×3.0×0.8mm），通过WLP "
        "（Wafer-Level Packaging）实现极致的微型化。",
    )

    heading_section(doc, "6.3  测量电路")

    heading_subsection(doc, "6.3.1  开关电容读出电路")
    body_paragraph(doc,
        "MEMS加速度计和陀螺仪的电容变化量极小（fF至aF量级），读出电路需要极高的灵敏度和分辨率。"
        "开关电容（Switched-Capacitor）读出电路是主流方案。其基本工作流程为：在驱动相（φ1），"
        "差分电容C1和C2被预充电至参考电压Vref；在检测相（φ2），电荷在C1和C2之间重新分布，"
        "产生与ΔC成正比的差分电压ΔV = Vref × (C1-C2)/(C1+C2) ≈ Vref × ΔC/(2C0)。该微弱差分电压"
        "经仪表放大器（INA, 增益100-1000倍）放大后，由Σ-Δ ADC转换为高分辨率数字信号（典型16-24bit）。"
        "Σ-Δ ADC利用过采样和噪声整形技术，将量化噪声推向高频段，在低频信号带宽内（100-500Hz）实现"
        ">100dB的信噪比。为抑制1/f闪烁噪声，读出电路常采用chopper stabilization技术，将信号调制至"
        "高频后再解调回基带。",
    )

    add_circuit_img(doc, "mems_readout_circuit", width_inches=5.2)

    heading_subsection(doc, "6.3.2  Allan方差分析与噪声辨识")
    body_paragraph(doc,
        "MEMS-IMU的误差源主要包括四类。零偏（Bias）指在无输入运动时传感器输出的非零均值，表现为陀螺仪零偏"
        "不稳定性（Bias Instability, 单位°/hr）和加速度计零偏（单位μg）。比例因子误差（Scale Factor Error）"
        "是传感器输出的斜率非线性，通常以ppm表示。交叉轴灵敏度（Cross-Axis Sensitivity）是各敏感轴之间的"
        "耦合误差，由微机械结构的非正交性和封装应力引入。角度随机游走（Angle Random Walk, ARW）是由热机械"
        "噪声（Brownian噪声）引起的白噪声积分，单位为°/√hr。"
    )
    body_paragraph(doc,
        "完整的IMU误差模型可表示为：a_meas = Sa × Ma × (a_true + ba + na) + wa，其中Sa为比例因子矩阵，"
        "Ma为轴失准矩阵，ba为加速度计零偏，na为加速度计随机噪声，wa为加速度计白噪声。陀螺仪的误差模型类似："
        "ω_meas = Sg × Mg × (ω_true + bg + ng) + wg。高精度标定（六面法、温度补偿）是减小系统误差的关键。"
    )

    body_paragraph(doc,
        "Allan方差分析是辨识MEMS-IMU各类噪声机制的经典方法。对静态采集的陀螺仪/加速度计数据，"
        "将数据分段并计算相邻段平均值的差值方差。Allan标准差σA(τ)对积分时间τ的双对数曲线中，"
        "不同斜率区域对应不同的噪声机制：斜率-1/2对应角度/速度白噪声（ARW/VRW），斜率0对应零偏"
        "不稳定性（1/f闪烁噪声），斜率+1/2对应角速率/加速度随机游走（RRW/RW），斜率+1对应量化"
        "噪声。Allan标准差在长积分时间处的极小值即为零偏不稳定性的定量指标。典型车载战术级"
        "MEMS-IMU的零偏不稳定性在0.5-5°/hr（陀螺仪）和5-50μg（加速度计）范围内。Allan方差"
        "分析为IMU选型和卡尔曼滤波噪声参数(Q矩阵)的设置提供了关键先验统计信息。"
    )

    add_formula_img(doc, "Allan", width_inches=4.5)

    heading_section(doc, "6.4  应用场景")
    body_paragraph(doc,
        "纯IMU的积分漂移在几十秒内就会使位置估计完全发散（漂移速率≈0.1-1°/s for MEMS）。因此，IMU在自动"
        "驾驶中从不单独使用，而是通过传感器融合与GPS/GNSS、视觉里程计（VO）和激光雷达里程计（LO）形成互补。"
    )
    body_paragraph(doc,
        "GPS/IMU组合导航：当GPS信号良好时（开阔道路），GPS提供绝对位置校正以抑制IMU漂移；当GPS信号中断时"
        "（隧道、城市峡谷），IMU提供短时高精度的航迹推算（Dead Reckoning）。视觉-惯性里程计（VIO）利用连续"
        "图像帧间的特征匹配联合IMU预积分进行紧耦合位姿估计。激光雷达-惯性里程计（LIO）利用点云配准与IMU预"
        "积分联合优化，如LIO-SAM框架，是当前高精地图构建和自动驾驶定位的主流方案，在KITTI、nuScenes等公开"
        "数据集上平移误差可小于0.5%。"
    )

    add_page_break(doc)

    # =====================================================================
    # 10. CHAPTER 7: 多传感器融合与深度学习
    # =====================================================================
    heading_chapter(doc, "第七章  多传感器融合与深度学习应用")

    heading_section(doc, "7.1  多传感器融合的必要性")
    body_paragraph(doc,
        "每一类传感器在工作原理、感知能力和失效模式上各有优劣，不存在单一种类的「全能」传感器。多传感器融合"
        "的根本目标是通过对异质传感器数据的时空对准与信息互补，实现任何单一传感器均无法达成的感知鲁棒性和"
        "准确性。下表从七个关键维度对五类传感器进行了综合对比。"
    )

    add_table(doc,
        headers=["性能维度", "摄像头", "毫米波雷达", "超声波雷达", "激光雷达", "IMU"],
        rows=[
            ["测距能力", "中（双目）", "优", "近距", "最优（cm级）", "无"],
            ["速度测量", "间接", "直接（多普勒）", "无", "直接（FMCW）", "间接（积分）"],
            ["角度分辨率", "优", "中（1-2°）", "差", "优（<0.1°）", "N/A"],
            ["纹理/颜色", "最优", "无", "无", "无（强度）", "无"],
            ["夜间/暗光", "差", "优", "优", "优（主动光源）", "优"],
            ["雨雾雪天", "差", "优", "中", "中-差", "优"],
            ["成本（单颗）", "$10-50", "$30-80", "$5-15", "$200-1000", "$10-50"],
        ],
        caption="表7-1  五大车载传感器综合性能对比",
        font_size=8
    )

    # Insert radar chart
    insert_image(doc, "radar_chart.png", width_inches=4.5,
                 caption="图7-1  五大车载传感器多维度性能雷达图对比")

    heading_section(doc, "7.2  三级融合架构")

    heading_subsection(doc, "7.2.1  数据级融合（Early Fusion）")
    body_paragraph(doc,
        "在原始数据或接近原始数据的层级进行融合。典型案例包括：将激光雷达点云投影到相机图像平面上，在每个"
        "像素点上附加LiDAR的深度和强度信息，从而生成RGB-D四通道图像；或将图像语义分割的结果反投影到点云上，"
        "为每个3D点赋予语义颜色标签。数据级融合最大限度地保留了原始信息的丰富性，但对传感器的时空标定精度"
        "要求极高——微小的外参标定误差（如0.1°的旋转偏差）在100m距离处即可引入约17cm的空间偏移，导致严重的"
        "特征失配。"
    )

    heading_subsection(doc, "7.2.2  特征级融合（Feature Fusion）")
    body_paragraph(doc,
        "在各传感器独立提取特征向量后进行融合。这一层级是目前自动驾驶感知系统的主流范式。代表性方法是BEV"
        "（鸟瞰图）特征融合：将摄像头、激光雷达和毫米波雷达各自提取的特征通过各自的投影/提升（Lift-Splat-"
        "Shoot）操作统一变换到BEV特征空间中，在BEV空间进行特征级联（Concatenation）或交叉注意力（Cross-"
        "Attention）。BEV融合的优势在于：(1) 统一坐标系消除了不同传感器视角带来的对应歧义；(2) BEV空间的"
        "平移不变性使CNN能够高效共享权重；(3) 融合后的特征可直接输入多任务头部网络。"
    )

    heading_subsection(doc, "7.2.3  决策级融合（Late Fusion）")
    body_paragraph(doc,
        "各传感器独立完成目标检测和分类，在目标级/轨迹级进行关联和融合。典型应用包括：各传感器各自输出目标"
        "列表（Object List），然后通过匈牙利算法或联合概率数据关联滤波器（JPDA）进行跨传感器的目标匹配和"
        "轨迹融合。决策级融合架构简单、传感器模块解耦、工程实现相对容易，但存在信息丢失的风险——若某传感器"
        "在其独立检测阶段漏检了目标，则该目标在融合阶段将彻底丢失，无法通过其他传感器的原始信息挽回。"
    )

    heading_section(doc, "7.3  BEV感知架构")
    body_paragraph(doc,
        "近年来，以Tesla的BEV感知架构和学术界提出的BEVFormer、BEVFusion等为代表，基于Transformer的鸟瞰图"
        "多模态融合已成为自动驾驶感知的主流算法范式。该架构的核心技术流程包括五个步骤。"
    )
    body_paragraph(doc,
        "第一步，多视图图像特征提取：使用ResNet/EfficientNet骨干网络从每个相机视角提取多尺度图像特征（C2-C5"
        "层），并通过FPN（特征金字塔网络）生成多尺度特征图。第二步，2D到BEV视图变换：通过Lift-Splat-Shoot"
        "（LSS）将每个像素提升（Lift）为沿视线方向的概率深度分布，然后将所有视锥点投影（Splat）到BEV网格中"
        "并求和池化。BEVFormer则使用可变形交叉注意力直接查询BEV网格点对应的多视图图像特征。"
    )
    body_paragraph(doc,
        "第三步，点云体素化与特征编码：将激光雷达点云通过体素化离散化后，利用PointPillars或VoxelNet进行稀疏"
        "特征提取，生成BEV点云特征。第四步，多模态特征融合：在BEV空间使用通道级联或可变形注意力机制，自适应"
        "地聚合来自不同传感器的BEV特征图。第五步，多任务检测头：融合后的BEV特征图输入多任务头部网络，同时"
        "输出3D目标检测、车道线分割、可行驶区域推理和轨迹预测等结果。"
    )

    heading_section(doc, "7.4  PointPillars点云目标检测")
    body_paragraph(doc,
        "PointPillars是工业界（NVIDIA/nutonomy，2019年CVPR）广泛采用的一种高效点云目标检测架构，巧妙地将"
        "稀疏无序的3D点云转化为适合2D卷积神经网络处理的伪图像表示。第一步，柱状化（Pillarization）：将三维"
        "空间在x-y（俯视图）平面上划分为等间距的网格（典型间距0.16m-0.2m），每个网格在z轴方向无限延伸形成"
        "一个「柱子（Pillar）」。落入同一柱子内的所有激光点构成该柱子的原始特征集。"
    )
    body_paragraph(doc,
        "第二步，PointNet特征学习：对每个非空柱子，利用简化版的PointNet将其中所有点的坐标和强度特征编码为"
        "一个固定长度的特征向量（通常64维），生成一个(C×H×W)的三维伪图像张量。第三步，2D卷积骨干网络：将"
        "伪图像输入到标准的2D CNN骨干网（如ResNet-18/34 + FPN/SSD检测头组合）进行多尺度特征提取和目标回归。"
        "PointPillars在KITTI 3D目标检测基准上实现了85%以上的mAP（Moderate难度），推理速度可达42FPS"
        "（NVIDIA Titan V），较好地平衡了精度与实时性，已成为工业界点云感知的事实标准之一。"
    )

    heading_section(doc, "7.5  TensorRT车载边缘部署")
    body_paragraph(doc,
        "自动驾驶对感知系统的端到端延迟有严格约束（通常要求<100ms）。深度学习模型在车载计算平台（如NVIDIA "
        "Jetson Orin、地平线征程5/6）上的高效部署涉及以下关键技术。模型量化（Model Quantization）将模型权重"
        "和激活值从FP32精度量化至FP16或INT8精度。经过校准后的INT8量化通常可在精度损失<1%的前提下，将推理"
        "速度提升2-4倍并降低约一半的内存带宽需求。"
    )
    body_paragraph(doc,
        "计算图层融合（Layer Fusion）通过TensorRT推理引擎自动识别并合并连续的操作层——如Conv+BN+ReLU合并为"
        "单一计算内核，消除中间张量的显存读写开销，减少GPU Kernel Launch次数。对于PointPillars等包含大量"
        "小算子的模型，层融合的加速效果尤为显著，可将吞吐量提升30%-50%。CUDA Kernel定制优化针对点云体素化"
        "中的大量不规则内存访问和原子加操作，工程师通过编写定制CUDA Kernel，利用共享内存作为用户管理缓存，"
        "有效降低Global Memory带宽压力。在Jetson Orin上，优化后的BEVFusion可实现25FPS@15W的端到端推理性能。"
    )

    add_page_break(doc)

    # =====================================================================
    # 11. CHAPTER 8: 市场前景与职业规划
    # =====================================================================
    heading_chapter(doc, "第八章  市场前景与职业规划")

    heading_section(doc, "8.1  传感器产业的关键技术趋势")
    body_paragraph(doc,
        "自动驾驶传感器产业当前正经历五个维度的深刻技术变革。"
    )
    body_paragraph(doc,
        "第一，固态化与芯片化。激光雷达正从混合固态（MEMS/转镜）向全固态（Flash/OPA）演进，硅光子集成使"
        "「LiDAR-on-a-Chip」的目标逐步接近。毫米波雷达从SiGe分立方案向CMOS RF-SoC单芯片方案过渡，单芯片集成"
        "度从收发前端延伸到DSP基带处理。超声波传感器从分立变压器方案向ASIC集成驱动方案转变。"
    )
    body_paragraph(doc,
        "第二，4D成像毫米波雷达商用化。增加俯仰维度的4D毫米波雷达于2024-2025年实现大规模前装量产，以<$100"
        "美元的单颗成本填补了传统毫米波雷达与激光雷达之间的性能鸿沟。4D雷达可输出超过10万点/帧的密集点云，"
        "有效检测静止障碍物、自由空间边界（路沿/护栏）和桥梁/立交桥。"
    )
    body_paragraph(doc,
        "第三，FMCW激光雷达。多家企业（Aeva、Mobileye、Scantinel等）计划在2026-2027年推出FMCW工程样机，"
        "以同时获取距离和瞬时速度的优势在L4 Robotaxi场景打开市场。FMCW的芯片级集成（PIC光子集成电路）是"
        "降低成本和体积的关键技术路径。"
    )
    body_paragraph(doc,
        "第四，AI与传感器深度融合。BEV+Transformer多模态感知架构在2023-2024年已成为技术主流，到2027年中国"
        "自动驾驶感知系统AI算法渗透率预计达92%。端到端自动驾驶（感知-预测-规划一体化神经网络）将成为下一代"
        "技术范式，对传感器的信息密度和质量提出更高要求。"
    )
    body_paragraph(doc,
        "第五，国产替代加速。激光雷达和毫米波雷达核心芯片的国产化率从2021年的不足15%跃升至2024年的约45%，"
        "2030年有望突破80%。在车载CIS、MMIC、VCSEL、SPAD和MEMS-IMU等核心元器件领域，中国厂商已初步形成完整"
        "的产业链布局。"
    )

    heading_section(doc, "8.2  成本下降曲线")
    body_paragraph(doc,
        "传感器模组的持续降价是推动自动驾驶从高端选配向主流车型渗透的核心经济驱动力。激光雷达模组（ADAS级）"
        "从2020年约10,000美元降至2024年约300-500美元，预计2027年<$200美元。4D毫米波雷达从2024年约80-120"
        "美元降至2027年预计<$50美元。L2+级全传感器套件（1LiDAR+5Radar+12Camera+1IMU）从2024年约8000-10000"
        "元降至2027年预计3000-4000元，2028-2030年有望进一步降至2000元以内，实现15万元级别经济型乘用车的标配。"
    )

    # Insert cost curve
    insert_image(doc, "cost_curve.png", width_inches=5.0,
                 caption="图8-1  自动驾驶传感器成本下降轨迹（2016-2030）")

    # Insert market players
    insert_image(doc, "market_players.png", width_inches=5.0,
                 caption="图8-2  2024年全球车载激光雷达市场份额（按供应商）")

    key_insight_box(doc,
        "自动驾驶传感器领域当前处于严重的人才供不应求状态——供给端培养速度远落后于需求端爆发速度。"
        "对电子信息类、计算机类、自动化类毕业生而言，传感器技术与自动驾驶感知融合是当前最具成长性的"
        "细分方向之一。顶尖企业（禾赛、速腾聚创、华为车BU）对应届硕士算法岗开出25-45万年薪，"
        "资深专家级工程师年薪可达80-150万元+。"
    )

    heading_section(doc, "8.3  五大典型岗位分析")
    body_paragraph(doc,
        "基于对禾赛科技、速腾聚创、其域创新、华为车BU、大疆车载、德赛西威等产业链核心企业的招聘信息分析，"
        "自动驾驶传感器领域的五大典型技术岗位及技能需求如下。"
    )

    body_paragraph(doc,
        "第一，激光雷达硬件工程师。核心技能包括：模拟/混合信号IC设计、光电探测器特性分析、TIA/TDC电路设计、"
        "SIP封装设计、VCSEL驱动电路设计。工具链为Cadence Virtuoso、ADS、HFSS、Altium Designer。知识背景"
        "要求光电信息科学与工程、微电子科学与工程、电子科学与技术。该岗位是传感器企业研发体系的核心，负责"
        "从芯片级到模组级的硬件架构设计，对物理电子学和模拟IC设计的理论基础要求极高。"
    )
    body_paragraph(doc,
        "第二，多传感器标定算法工程师。核心技能包括：多视图几何、非线性优化（Ceres/g2o）、PnP/手眼标定、"
        "LiDAR-Camera联合标定、C++/Python编程。知识背景要求计算机视觉、机器人学、摄影测量与遥感。标定是"
        "传感器融合的基础——标定精度直接决定了融合算法的性能上限。传感器之间的外参标定（包括旋转矩阵R和平移"
        "向量t）通常通过最小化重投影误差或点面距离的束调整（Bundle Adjustment）优化来完成。"
    )
    body_paragraph(doc,
        "第三，感知融合算法工程师。核心技能包括：3D目标检测（PointPillars/CenterPoint）、BEV感知/Transformer"
        "架构、多目标跟踪（MOT/AB3DMOT）、传感器前/后融合、PyTorch/TensorRT。知识背景要求计算机科学"
        "（人工智能方向）、模式识别与智能系统。该岗位是当前需求量最大、薪资最高的方向，需要深入理解深度"
        "学习模型的设计、训练与部署全流程。"
    )
    body_paragraph(doc,
        "第四，自动驾驶嵌入式系统工程师。核心技能包括：GPGPU/CUDA编程、模型量化与TensorRT部署、QNX/Linux"
        "实时操作系统、DDS/SomeIP中间件、ROS2。知识背景要求计算机体系结构、嵌入式系统、软件工程。该岗位"
        "连接算法与硬件，是决定感知系统实时性能的关键角色。"
    )
    body_paragraph(doc,
        "第五，传感器系统集成与测试工程师。核心技能包括：传感器性能测试（MTF、角分辨率、探测概率）、EMC/EMI"
        "测试、车规级可靠性验证（AEC-Q100/Q200）、HIL/SIL测试。知识背景要求车辆工程、测控技术与仪器、电气"
        "工程。该岗位确保传感器满足车载环境严苛的可靠性要求（-40~+105°C温度范围、高湿度、振动冲击等）。"
    )

    heading_section(doc, "8.4  薪资水平与发展前景")

    add_table(doc,
        headers=["岗位类别", "应届硕士", "3-5年经验", "资深/专家(5-10年+)"],
        rows=[
            ["感知融合算法(AI/深度学习)", "28-45万", "50-85万", "90-160万+"],
            ["标定与定位算法", "25-40万", "45-75万", "80-140万+"],
            ["激光雷达/毫米波硬件设计", "22-35万", "40-65万", "65-120万+"],
            ["嵌入式系统/车载软件", "20-32万", "35-55万", "55-100万+"],
            ["系统集成与测试验证", "15-28万", "28-50万", "45-85万+"],
        ],
        caption="表8-1  自动驾驶传感器领域典型岗位薪资区间（2024-2025年，年薪万元）",
        font_size=9
    )

    body_paragraph(doc,
        "关键行业趋势："
    )
    body_paragraph(doc,
        "第一，人才缺口持续扩大。中国智能网联汽车产业2025年直接人才缺口约35-50万人，其中传感器感知与算法"
        "方向占比约30%，缺口规模约10-15万人。"
    )
    body_paragraph(doc,
        "第二，薪资增速显著高于传统制造业。自动驾驶传感器领域应届硕士薪资较传统电子信息制造业高50%-120%，"
        "3-5年经验后差距进一步扩大至2-3倍。"
    )
    body_paragraph(doc,
        "第三，地域集中度高。上海、深圳、北京、苏州、广州五城市集中了约80%的自动驾驶传感器相关岗位。"
        "广深地区因比亚迪、小鹏、华为车BU、大疆车载、德赛西威、其域创新等企业的集聚效应，近年来岗位"
        "增速最为迅猛。"
    )
    body_paragraph(doc,
        "第四，学历与技能门槛。算法类岗位普遍要求硕士及以上学历（约78%的感知算法岗位要求硕士学历）；"
        "硬件和嵌入式类岗位对硕士学历的硬性要求相对低（约40%-55%），但更看重项目实践经验和芯片/板级调试能力。"
    )
    body_paragraph(doc,
        "该领域目前处于人才供需严重失衡的状态——供给端培养速度远落后于需求端爆发速度。对电子信息类"
        "（含光电信息、微电子）、计算机类（含AI、软件工程）、自动化类（含模式识别、测控技术）专业的毕业生"
        "而言，传感器技术与自动驾驶感知融合是当前最具成长性的细分方向之一。建议在校同学通过以下路径提升"
        "竞争力：(1) 参与开源自动驾驶项目（如Apollo、Autoware）并积累真实传感器数据处理经验；(2) 在"
        "KITTI/nuScenes/Waymo Open Dataset等公开数据集上复现主流感知模型；(3) 掌握至少一种嵌入式平台"
        "的模型部署工具链（TensorRT/NCNN/ONNX Runtime）；(4) 关注目标企业（禾赛、速腾聚创、其域创新、"
        "华为车BU、大疆车载等）的实习招聘周期，提前准备。"
    )

    add_page_break(doc)

    # =====================================================================
    # 12. CHAPTER 9: 结论与展望
    # =====================================================================
    heading_chapter(doc, "第九章  结论与展望")

    body_paragraph(doc,
        "本文对自动驾驶汽车中的五类先进传感器——视觉摄像机、毫米波雷达、超声波雷达、激光雷达和惯性测量单元"
        "——从物理原理、核心元器件、测量电路到多传感器融合应用进行了系统性综述。"
    )
    body_paragraph(doc,
        "在物理原理层面，本文以内光电效应、电磁波FMCW信号模型、压电换能效应和牛顿惯性定律为物理基础，建立"
        "了各传感器从原始物理量到可测量电信号的完整数学模型。从CMOS图像传感器的Iph = R × Popt光电转换方程，"
        "到FMCW毫米波雷达的fIF = (2B×R)/(c×Tc)距离-中频映射，从超声波雷达的R = vsΔt/2渡越时间公式，到激光"
        "雷达ToF的R = cΔt/2和FMCW相干探测模型，以及MEMS陀螺仪的Fc = -2m(Ω×v)科里奥利力方程——这些数学"
        "模型共同构成了车载传感器的物理学理论基础。"
    )
    body_paragraph(doc,
        "在核心元器件层面，从CMOS有源像素（4T-APS pinned光电二极管结构）到VCSEL/SPAD光电收发对，从77GHz "
        "CMOS MMIC（集成PLL+PA+LNA+Mixer+ADC）到MEMS振动陀螺仪（驱动-检测双模态谐振器），详细剖析了各器件"
        "的微观工作机制和关键技术指标。特别地，在激光雷达接收链路的分析中，从PIN→APD→SPAD→SiPM的探测器"
        "灵敏度演进路径揭示了单光子探测技术的工程实现细节。"
    )
    body_paragraph(doc,
        "在测量电路层面，以CDS相关双采样（Vout = Vsignal - Vreset）、TIA跨阻放大器（H(s) = -Rf/(1+sRfCf)）"
        "和TDC时间数字转换器（Δt_fine = N × τLSB, τLSB = 10-20ps）为代表性案例，深入阐述了从微弱的模拟"
        "物理量（fA-pA级电流、μV-mV级电压、ns-ps级时间间隔）到高精度数字量输出的信号调理与转换链路。这些"
        "测量电路的性能直接决定了传感器的精度、灵敏度和动态范围。"
    )
    body_paragraph(doc,
        "在应用与展望方面，本文指出：多传感器融合是自动驾驶感知系统不可逆转的技术方向，BEV+Transformer的"
        "前融合架构已成为产业主流范式，而GPGPU/TensorRT驱动的模型轻量化部署将深度学习感知推向了车载实时推理"
        "的边缘（25FPS@15W on Jetson Orin）。市场数据表明，2025-2030年是中国自动驾驶传感器产业从「技术验证」"
        "迈向「规模化商业落地」的黄金窗口期，激光雷达和4D成像毫米波雷达是增长确定性最高的两大细分赛道。"
    )
    body_paragraph(doc,
        "未来，随着硅光子学、3D堆叠封装（3D-Stacking/TSV）和存算一体（PIM/Compute-in-Memory）架构的持续"
        "推进，「传感-计算-决策」三者将在物理芯片层面实现深度融合，形成传感计算一体化（Sensing-Computing "
        "Convergence）的全新范式。对于有志于从事传感器技术与智能驾驶交叉领域的学生而言，扎实的物理电子学"
        "基础、电机拖动与控制理论、模拟/混合信号IC设计与嵌入式系统实践、深度学习算法能力（PyTorch/TensorRT"
        "生态），共同构成了进入这一前沿行业的综合竞争力。"
    )

    add_page_break(doc)

    # =====================================================================
    # 13. REFERENCES
    # =====================================================================
    heading_chapter(doc, "参考文献")

    references = [
        "[1] 王庞伟等. 智能网联汽车电子技术[M]. 北京: 机械工业出版社, 2021.",
        "[2] 周润景等. 常用传感器技术及应用（第二版）[M]. 北京: 电子工业出版社, 2021.",
        "[3] 李立功等. 光电传感器原理与应用[M]. 北京: 科学出版社, 2020.",
        "[4] 程润伟等. CUDA C编程权威指南[M]. 北京: 机械工业出版社, 2017.",
        "[5] 景乃锋等. 通用图形处理器设计：GPGPU编程模型与架构原理[M]. 北京: 清华大学出版社, 2022.",
        "[6] 方吉鑫, 等. 硅光电倍增管前端放大电路的架构研究进展[J]. 激光技术, 2026, 50(2): 161-168.",
        "[7] 林远启, 等. 集成化多线列激光雷达模拟前端微组件设计[J]. 光电工程, 2021, 48(2): 210080.",
        "[8] 王伟, 张辉, 李明. 车载固态激光雷达技术发展综述[J]. 红外与激光工程, 2024, 53(1): 20230330.",
        "[9] 刘伟奇. Review of Automotive Sensors Based on Autonomous Driving[C]. IEEE Sensors, 2025.",
        "[10] Lang A H, Vora S, Caesar H, et al. PointPillars: Fast encoders for object detection "
              "from point clouds[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and "
              "Pattern Recognition (CVPR). 2019: 12697-12705.",
        "[11] Li Y, Ibanez-Guzman J. Lidar for autonomous driving: The principles, challenges, and "
              "trends for automotive lidar and perception systems[J]. IEEE Signal Processing Magazine, "
              "2020, 37(4): 50-61.",
        "[12] Behroozpour B, Sandborn P A M, Wu M C, et al. Lidar system architectures and circuits[J]. "
              "IEEE Communications Magazine, 2017, 55(10): 135-142.",
        "[13] Poulton C V, et al. Optical beamforming techniques for solid-state automotive LIDAR[R]. "
              "UC Berkeley EECS Technical Report, 2023.",
        "[14] Yole Intelligence. LiDAR for Automotive and Industrial Applications 2024[R]. Yole Group, 2024.",
        "[15] 高工智能汽车研究院. 2024年度中国乘用车前装激光雷达市场报告[R]. 2025.",
        "[16] 中商产业研究院. 2024-2029年中国激光雷达行业市场深度调研及发展前景报告[R]. 2024.",
        "[17] 中商产业研究院. 全球激光雷达市场规模报告[R]. 2024.",
        "[18] QYResearch. 全球机器人激光雷达市场预测（2023-2030）[R]. 2024.",
        "[19] 禾赛科技. AT128产品技术规格书[Z]. 2023.",
        "[20] 速腾聚创. RS-LiDAR-M1/E1平台技术白皮书[Z]. 2024.",
        "[21] 加特兰微电子. Alps系列77GHz CMOS毫米波雷达SoC数据手册[Z]. 2024.",
        "[22] 东兴证券研究所. 激光雷达产业链深度分析报告[R]. 2024.",
        "[23] 其域创新. 产品与招聘信息[EB/OL]. https://xgrids.cn. 2025.",
        "[24] 其域创新. 多传感器标定算法工程师招聘[EB/OL]. "
              "https://pecivkvtit.jobs.feishu.cn/index/position/7486766416390113587/detail. 2025.",
        "[25] 其域创新. 感知融合算法工程师招聘[EB/OL]. "
              "https://pecivkvtit.jobs.feishu.cn/index/position/7491253117397387529/detail. 2025.",
        "[26] 其域创新. 多传感器融合算法工程师(视觉雷达联合重建方向)招聘[EB/OL]. "
              "https://pecivkvtit.jobs.feishu.cn/index/position/7535787737665538350/detail. 2025.",
        "[27] IEEE Standard for Inertial Sensor Terminology. IEEE Std 528-2019[S]. 2019.",
        "[28] El-Sheimy N, Hou H, Niu X. Analysis and modeling of inertial sensors using Allan "
              "variance[J]. IEEE Transactions on Instrumentation and Measurement, 2008, 57(1): 140-149.",
        "[29] Bronzi D, Villa F, Tisa S, et al. 100 000 frames/s 64x32 single-photon detector array "
              "for 2-D imaging and 3-D ranging[J]. IEEE Journal of Selected Topics in Quantum "
              "Electronics, 2014, 20(6): 354-363.",
        "[30] Pellegrini S, Rae B, Ping G A. SPAD sensors: from single photon detection to LiDAR[J]. "
              "Nature Photonics, 2022, 16: 441-448.",
        "[31] Zappa F, Tisa S, Tosi A, et al. Principles and features of single-photon avalanche "
              "diode arrays[J]. Sensors and Actuators A: Physical, 2007, 140(1): 103-112.",
        "[32] Cova S, Ghioni M, Lacaita A, et al. Avalanche photodiodes and quenching circuits "
              "for single-photon detection[J]. Applied Optics, 1996, 35(12): 1956-1976.",
        "[33] 华为技术有限公司. 华为4D成像毫米波雷达技术白皮书[Z]. 2024.",
        "[34] Continental AG. ARS540 4D Imaging Radar Technical Data Sheet[Z]. 2023.",
        "[35] 德赛西威. 自动驾驶域控制器与传感器集成方案[Z]. 2024.",
    ]

    for ref in references:
        p = new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=1, space_after=1,
                          first_line_indent=None, line_spacing=1.3)
        add_run(p, ref, 'SimSun', 'SimSun', 12, bold=False)

    add_page_break(doc)

    # =====================================================================
    # 14. ACKNOWLEDGMENTS
    # =====================================================================
    heading_chapter(doc, "致　谢")

    body_paragraph(doc,
        "感谢任课教师翟老师（zhaiyh@gdut.edu.cn）在《传感器技术与应用》课程中的系统讲授与悉心指导。"
        "课程中关于光电传感器、CCD/CMOS图像传感器、内光电效应、暗电流机理以及信号调理电路的深入讲解，"
        "为本调研报告的撰写奠定了扎实的理论基础。翟老师在授课中将理论公式推导与工程实践案例有机结合，"
        "使学生在理解传感器物理原理的同时，对实际电路设计中的性能约束与工程权衡有了直观的认知，这种"
        "\「原理—器件—电路\」三位一体的讲授方法对本文写作产生了深刻影响。"
    )
    body_paragraph(doc,
        "感谢广东工业大学提供的丰富的数字图书馆资源（包括IEEE Xplore、中国知网CNKI、万方数据库等），"
        "使本报告得以广泛检索和引用来自光电工程、激光技术、IEEE Signal Processing Magazine、CVPR等"
        "高水平期刊与会议的学术文献。"
    )
    body_paragraph(doc,
        "本报告写作过程中参考和引用了大量来自国内外激光雷达企业（禾赛科技、速腾聚创、其域创新、"
        "加特兰微电子等）、学术研究机构（UC Berkeley、MIT、清华大学、中国科学院等）以及市场研究机构"
        "（Yole Intelligence、高工智能汽车研究院、QYResearch、中商产业研究院等）的公开技术文献和市场数据。"
        "禾赛科技AT128、速腾聚创E1平台的技术文档为第五章激光雷达的器件分析提供了宝贵的一手工程资料；"
        "加特兰微电子的Alps系列芯片数据手册为第三章毫米波雷达CMOS工艺国产替代趋势的判断提供了关键依据。"
        "在此向这些推动传感器技术进步的组织和个人致以诚挚谢意。"
    )
    body_paragraph(doc,
        "最后，感谢家人和朋友们在调研写作期间的理解与支持。传感器的世界从微观的皮秒计时到宏观的市场"
        "千亿规模，跨度之大令人在调研中既感敬畏又觉兴奋。希望这份报告不仅能完成课程考核要求，也能为"
        "同样对自动驾驶传感器技术感兴趣的同学提供一份有参考价值的入门材料。"
    )

    # ── Save ─────────────────────────────────────────────────────────
    output_path = r"D:\sensorhomework\论文\自动驾驶汽车先进传感器技术综述.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


# ── Main ─────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    create_document()
